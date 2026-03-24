# import os
# import fitz
# import openai
# import pandas as pd
# from dotenv import load_dotenv
# from docx import Document
# import logging
# import re
# import easyocr
# from PIL import Image
# import io
# import numpy as np
# import json

# # === Setup Logging ===
# logging.basicConfig(
#     filename='funding_nc_analyzer.log',
#     level=logging.DEBUG,
#     format='%(asctime)s - %(levelname)s - %(message)s'
# )

# # === Load environment variables ===
# load_dotenv()
# openai.api_key = os.getenv("OPENAI_API_KEY")

# # === File Paths ===
# JSON_DATA_PATH = "data/card_data.json"
# ENRICH_PDF_FILES = [
#     "data/General Credit Card Knowledge.pdf",
#     "data/DATA POINTS - BUSINESS CREDIT CARD DATA POINTS.pdf",
#     "data/HOW To Leverage Business Credit to.pdf"
# ]
# ENRICH_DOCX_FILES = [
#     "data/Credit Stacking Guide Lines and Better Practices.docx"
# ]
# ENRICH_CSV_FILES = [
#     "data/Tarjetas de Negocio sin Garantia Personal.csv"
# ]

# # Initialize EasyOCR reader
# reader = easyocr.Reader(['en'], gpu=False)

# # === Function to load JSON data ===
# def load_json_data():
#     try:
#         with open(JSON_DATA_PATH, encoding="utf-8") as file:
#             data = json.load(file)
#             logging.info(f"Loaded JSON data from {JSON_DATA_PATH}")
#             print(f"✅ Loaded JSON data with {len(data)} cards")
#             return data
#     except Exception as e:
#         logging.error(f"Failed to load JSON data: {str(e)}")
#         print(f"❌ Failed to load JSON data: {str(e)}")
#         return None

# # === Function to extract text from PDF ===
# def extract_text_from_pdf(pdf_path):
#     try:
#         text = ""
#         with fitz.open(pdf_path) as doc:
#             for page in doc:
#                 page_text = page.get_text()
#                 if page_text.strip():
#                     text += page_text + "\n"
#                 else:
#                     logging.info(f"No text found in page {page.number + 1} for {pdf_path}, attempting OCR")
#                     pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
#                     img_bytes = pix.tobytes("png")
#                     img = Image.open(io.BytesIO(img_bytes))
#                     img_np = np.array(img)
#                     ocr_result = reader.readtext(img_np, detail=0)
#                     ocr_text = "\n".join(ocr_result)
#                     if ocr_text.strip():
#                         text += ocr_text + "\n"
#                     else:
#                         logging.warning(f"No text extracted via OCR from page {page.number + 1} of {pdf_path}")
        
#         if text.strip():
#             logging.info(f"Successfully extracted text from PDF: {pdf_path}")
#             logging.debug(f"Extracted text snippet: {text[:500]}...")
#             return text
#         else:
#             logging.error(f"No text could be extracted from {pdf_path}, even with OCR")
#             return None
#     except Exception as e:
#         logging.error(f"Error extracting PDF text from {pdf_path}: {str(e)}")
#         return None

# # === New Function to extract text from JSON input ===
# def extract_text_from_json(json_input):
#     try:
#         data = json.loads(json_input)
#         text = ""
#         # Assuming JSON structure like: {"credit_score": 750, "utilization": "5%", "history": "Clean", ...}
#         # Convert JSON to text format similar to PDF extraction for analysis
#         for key, value in data.items():
#             text += f"{key.capitalize()}: {value}\n"
#         if text.strip():
#             logging.info("Successfully extracted text from JSON input")
#             logging.debug(f"Extracted text snippet: {text[:500]}...")
#             return text
#         else:
#             logging.error("No data found in JSON input")
#             return None
#     except Exception as e:
#         logging.error(f"Error extracting text from JSON input: {str(e)}")
#         return None
    
# def extract_text_from_json(json_input):
#     try:
#         data = json.loads(json_input)
#         text = ""
#         # Assuming JSON structure like: {"credit_score": 750, "utilization": "5%", "history": "Clean", ...}
#         # Convert JSON to text format similar to PDF extraction for analysis
#         for key, value in data.items():
#             text += f"{key.capitalize()}: {value}\n"
#         if text.strip():
#             logging.info("Successfully extracted text from JSON input")
#             logging.debug(f"Extracted text snippet: {text[:500]}...")
#             return text
#         else:
#             logging.error("No data found in JSON input")
#             return None
#     except Exception as e:
#         logging.error(f"Error extracting text from JSON input: {str(e)}")
#         return None

# # === Function to extract Credit Score and Utilization ===
# def extract_credit_info(text):
#     score_pattern = r"Credit Score\s*[:\-]?\s*(\d{3,4})"
#     score_matches = re.findall(score_pattern, text)
#     score = score_matches[0] if score_matches else None
    
#     utilization_pattern = r"Utilization\s*[:\-]?\s*(\d{1,3}%?)"
#     utilization_matches = re.findall(utilization_pattern, text)
#     utilization = utilization_matches[0] if utilization_matches else None
    
#     return score, utilization

# # === Function to extract text from DOCX ===
# def extract_text_from_docx(docx_path):
#     try:
#         doc = Document(docx_path)
#         text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
#         logging.info(f"Successfully extracted text from DOCX: {docx_path}")
#         return text
#     except Exception as e:
#         logging.error(f"Error extracting DOCX text from {docx_path}: {str(e)}")
#         return None

# # === Function to extract text from CSV ===
# def extract_text_from_csv(csv_path):
#     try:
#         df = pd.read_csv(csv_path)
#         text = df.to_string(index=False)
#         return text
#     except Exception as e:
#         logging.error(f"Error extracting CSV text from {csv_path}: {str(e)}")
#         return None

# # === Support Functions ===
# def get_state_funding_cards(user_state, json_data):
#     try:
#         state_cards = []
#         for card_name, card_info in json_data.items():
#             if user_state in card_info['state']:
#                 state_cards.append({
#                     'card_name': card_name,
#                     'bank': card_info['bank'],
#                     'apr': card_info['apr'],
#                     'mode': card_info['mode'],
#                     'bureau': card_info['bureau'],
#                     'is_travel': card_info['is_travel'],
#                     'game_plan': card_info['game_plan'],
#                     'tips_info': card_info.get('tips_info', 'No additional tips available for this card.')
#                 })
#         if not state_cards:
#             logging.error(f"No cards found for state {user_state}")
#             print(f"❌ No cards found for state {user_state}")
#             return []
#         logging.info(f"Found {len(state_cards)} cards for state {user_state}: {[card['card_name'] for card in state_cards]}")
#         print(f"✅ Found {len(state_cards)} cards for state {user_state}: {[card['card_name'] for card in state_cards]}")
#         return state_cards
#     except Exception as e:
#         logging.error(f"Error loading state funding cards: {str(e)}")
#         print(f"❌ Error loading state funding cards: {str(e)}")
#         return []

# def get_enrichment():
#     enrichment = ""
#     for file in ENRICH_PDF_FILES + ENRICH_DOCX_FILES + ENRICH_CSV_FILES:
#         if os.path.exists(file):
#             try:
#                 if file.endswith(".pdf"):
#                     text = extract_text_from_pdf(file)
#                 elif file.endswith(".docx"):
#                     text = extract_text_from_docx(file)
#                 elif file.endswith(".csv"):
#                     text = extract_text_from_csv(file)
#                 else:
#                     text = ""
#                 if text:
#                     enrichment += f"\n[From {os.path.basename(file)}]\n{text[:1000]}...\n"
#             except Exception as e:
#                 enrichment += f"\n[Error reading {file}]: {str(e)}\n"
#                 logging.error(f"Error reading enrichment file {file}: {str(e)}")
#         else:
#             enrichment += f"\n[Skipped missing file: {file}]\n"
#             logging.warning(f"Skipped missing enrichment file: {file}")
#     return enrichment

# # === Output Validation ===
# def validate_gpt_output(analysis, state_cards, user_state, json_data, mode="free"):
#     """
#     Validates and updates the GPT-4 output to ensure all required sections are present,
#     cards are valid for the user's state, travel cards are included in specified order,
#     non-travel cards are from unique banks, and correct bureaus are assigned.
    
#     Args:
#         analysis (str): The raw GPT-4 output string.
#         state_cards (list): List of state-specific card dictionaries from JSON data.
#         user_state (str): The user's selected U.S. state.
#         json_data (dict): The JSON data containing card information.
#         mode (str): User mode ('free' or 'paid').
    
#     Returns:
#         str: Updated analysis string with validated sections and complete game_plan/tips_info in table.
#     """
#     logging.debug(f"Validating GPT-4 output for cards in state {user_state}, mode: {mode}")
    
#     analysis_lower = analysis.lower()
#     not_qualified = "does not qualify for funding" in analysis_lower
#     eligible_message = "you are eligible for funding" in analysis_lower

#     # Remove any mention of 'Inferred' or variations
#     analysis = re.sub(r"\(Inferred\)", "", analysis, flags=re.IGNORECASE)
#     analysis = re.sub(r"inferred as \d+\b", lambda m: m.group(0).replace("inferred as ", ""), analysis, flags=re.IGNORECASE)
#     analysis = re.sub(r"\bInferred\b", "", analysis, flags=re.IGNORECASE)

#     # Handle 'Data not available' replacements
#     if "Data not available" in analysis:
#         logging.warning("GPT used 'Data not available' in output. Replacing with estimated values.")
#         replacements = {
#             r"Credit Score.*Data not available": "Credit Score: 700 (based on industry standards and clean payment history)",
#             r"Utilization.*Data not available": "Utilization: 15% (based on typical credit profiles with high-limit cards)",
#             r"Avg\. Credit Age.*Data not available": "Avg. Credit Age: 2.5 years (based on standard account age)",
#             r"Hard Inquiries.*Data not available": "Hard Inquiries: 2 (based on typical inquiry patterns)"
#         }
#         for pattern, replacement in replacements.items():
#             analysis = re.sub(pattern, replacement, analysis, flags=re.IGNORECASE)
#         analysis += "\n\n📋 Note: Some values were estimated based on industry standards to provide a complete analysis."

#     # Free mode handling
#     if mode == "free":
#         logging.info("Validating and cleaning GPT output for free mode.")
#         if eligible_message:
#             logging.info("User is eligible in free mode. Ensuring Section 5 and Section 7 reflect eligibility.")
#             # Update Section 5 (Verdict)
#             analysis = re.sub(
#                 r"📌 \*\*5\. Verdict\*\*.*?(?=\n\n📌 \*\*6\. Action Plan\*\*|\Z)",
#                 f"📌 **5. Verdict**\n\n"
#                 f"🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.\n",
#                 analysis,
#                 flags=re.DOTALL
#             )
#             # Update Section 7 (Recommended Funding Sequence)
#             analysis = re.sub(
#                 r"📌 \*\*7\. Recommended Funding Sequence \((.*?)\)\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
#                 f"📌 **7. Recommended Funding Sequence ({user_state})**\n\n"
#                 f"🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.\n\n",
#                 analysis,
#                 flags=re.DOTALL
#             )
#         else:
#             logging.info("User is not qualified in free mode. Ensuring Section 5 and Section 7 reflect ineligibility.")
#             # Update Section 5 (Verdict)
#             analysis = re.sub(
#                 r"📌 \*\*5\. Verdict\*\*.*?(?=\n\n📌 \*\*6\. Action Plan\*\*|\Z)",
#                 f"📌 **5. Verdict**\n\n"
#                 f"Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.\n",
#                 analysis,
#                 flags=re.DOTALL
#             )
#             # Update Section 7 (Recommended Funding Sequence)
#             analysis = re.sub(
#                 r"📌 \*\*7\. Recommended Funding Sequence \((.*?)\)\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
#                 f"📌 **7. Recommended Funding Sequence ({user_state})**\n\n"
#                 f"Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.\n\n",
#                 analysis,
#                 flags=re.DOTALL
#             )
#         # Remove Game Plan and Tips and Info sections in free mode
#         analysis = re.sub(
#             r"\*\*Game Plan\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
#             "**Disclaimer**",
#             analysis,
#             flags=re.DOTALL
#         )
        
#         # Add contact message to the end of the output
#         contact_message = (
#             "\n\nIf you need more detailed guidance or have any questions, "
#             "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#         )
#         analysis += contact_message
#         logging.info("Added contact message to the final output for free mode.")
#         return analysis

#     # Paid mode handling
#     if mode == "paid":
#         if not_qualified:
#             logging.info("User is not qualified in paid mode. Updating Section 5 and Section 7 to reflect ineligibility and removing card recommendations.")
#             # Update Section 5 (Verdict)
#             analysis = re.sub(
#                 r"📌 \*\*5\. Verdict\*\*.*?(?=\n\n📌 \*\*6\. Action Plan\*\*|\Z)",
#                 f"📌 **5. Verdict**\n\n"
#                 f"Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.\n",
#                 analysis,
#                 flags=re.DOTALL
#             )
#             # Update Section 7 (Recommended Funding Sequence)
#             analysis = re.sub(
#                 r"📌 \*\*7\. Recommended Funding Sequence \((.*?)\)\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
#                 f"📌 **7. Recommended Funding Sequence ({user_state})**\n\n"
#                 f"Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.\n\n",
#                 analysis,
#                 flags=re.DOTALL
#             )
#             # Remove Game Plan and Tips and Info sections
#             analysis = re.sub(
#                 r"\*\*Game Plan\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
#                 "**Disclaimer**",
#                 analysis,
#                 flags=re.DOTALL
#             )
#             # Generate general guidance for Game Plan
#             new_insights = ""
#             if "Credit scores are below the required 720 threshold" in analysis:
#                 new_insights += "- **Improve Credit Score**: Focus on timely payments and reducing credit card balances to boost your credit score above 720.\n"
#             if "Util kvarization is above the ideal 10%" in analysis:
#                 new_insights += "- **Reduce Utilization**: Pay down credit card balances to achieve utilization below 10% to improve funding eligibility.\n"
#             if "inquiries are high" in analysis_lower:
#                 new_insights += "- **Limit Inquiries**: Avoid applying for new credit to keep inquiries low (≤ 3 in 6 months).\n"
#             new_insights += "- **Monitor Credit Profile**: Regularly check your credit reports for errors and dispute inaccuracies to strengthen your profile.\n"
#             new_insights += "- **Contact Negocio Capital**: Schedule a consultation for personalized guidance on improving your credit profile.\n"
            
#             analysis = re.sub(
#                 r"(\*\*Disclaimer\*\*|\Z)",
#                 f"**Game Plan**\n{new_insights}\n\1",
#                 analysis,
#                 flags=re.DOTALL
#             )
            
#             # Add contact message to the end of the output
#             contact_message = (
#                 "\n\nIf you need more detailed guidance or have any questions, "
#                 "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#             )
#             analysis += contact_message
#             logging.info("Added contact message to the final output for paid mode (not qualified).")
#             return analysis

#         if eligible_message:
#             logging.info("User is eligible in paid mode. Ensuring correct card recommendations with Game Plan and Tips and Info in table.")
#             if "please upgrade to our Premium Plan" in analysis_lower:
#                 logging.error("Incorrect verdict message for paid mode")
#                 analysis = re.sub(
#                     r"🎉 You're eligible for funding! To view your matched bank recommendations.*?Plan\.",
#                     "🎉 You're eligible for funding! See your matched bank recommendations below.",
#                     analysis
#                 )
#                 logging.info("Fixed incorrect verdict message for paid mode.")

#             # Define travel cards and validate availability
#             travel_cards = [
#                 'Chase Ink Unlimited',
#                 'BOFA Alaska Airlines Business',
#                 'Chase Sapphire Preferred'
#             ]
#             available_travel_cards = [card for card in travel_cards if any(c['card_name'] == card for c in state_cards)]
#             if len(available_travel_cards) < 3:
#                 logging.error(f"Only {len(available_travel_cards)} travel cards available for {user_state}. Expected 3: {', '.join(travel_cards)}.")
#                 error_note = f"\n\n⚠️ ERROR: Only {len(available_travel_cards)} travel cards available for {user_state}. Expected 3: {', '.join(travel_cards)}. Please select a different state or contact Negocio Capital for assistance."
#                 analysis += error_note
                
#                 # Add contact message to the end of the output
#                 contact_message = (
#                     "\n\nIf you need more detailed guidance or have any questions, "
#                     "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#                 )
#                 analysis += contact_message
#                 logging.info("Added contact message to the final output for paid mode (insufficient travel cards).")
#                 return analysis

#             non_travel_cards = [card['card_name'] for card in state_cards if card['card_name'] not in travel_cards]
#             json_card_names = {card['card_name'] for card in state_cards}

#             # Count unique banks
#             unique_banks = set(card['bank'] for card in state_cards)
#             num_unique_banks = len(unique_banks)
#             logging.info(f"Number of unique banks for {user_state}: {num_unique_banks}")

#             # Remove ROUND 3 if num_unique_banks < 6
#             num_rounds = 2 if num_unique_banks < 6 else 3
#             if num_unique_banks < 6:
#                 analysis = re.sub(
#                     r"\*\*ROUND 3\*\*.*?(?=\*\*Game Plan|\*\*Disclaimer\*\*|\Z)",
#                     "",
#                     analysis,
#                     flags=re.DOTALL
#                 )
#                 logging.info("Removed ROUND 3 from analysis due to insufficient unique banks (< 6).")

#             # Priority banks for non-travel cards
#             priority_banks = ['Chase', 'American Express', 'BMO Harris', 'KeyBank', 'Truist', 'Bank of America']
#             available_priority_banks = [bank for bank in priority_banks if any(card['bank'] == bank and card['card_name'] in non_travel_cards for card in state_cards)]

#             # Initialize tracking
#             non_travel_cards_selected = {}  # Track non-travel card banks
#             travel_cards_selected = {}     # Track travel card banks
#             valid_modes = ["Online", "In-branch", "Phone", "Online (requires account)", "Online (Omaha Zip)", "Phone/In-branch", "In-branch/Phone"]
#             all_card_names = []            # Track all selected cards for validation
#             chase_zero_apr_count = 0       # Track number of 0% APR Chase cards
#             bureau_usage = {'Experian': 0, 'TransUnion': 0, 'Equifax': 0}  # Track bureau usage for balancing

#             # Define bureau assignments for each round
#             bureau_assignments = {
#                 'ROUND 1': [
#                     ('Chase Ink Unlimited', 'Experian', 'Supports business rewards'),
#                     (None, 'TransUnion', None),
#                     (None, 'Equifax', None)
#                 ],
#                 'ROUND 2': [
#                     ('BOFA Alaska Airlines Business', 'TransUnion', 'Supports travel rewards'),
#                     (None, 'Experian', None),
#                     (None, 'Equifax', None)
#                 ],
#                 'ROUND 3': [
#                     (None, 'Equifax', None),
#                     ('Chase Sapphire Preferred', 'Experian', 'Supports travel rewards'),
#                     (None, 'TransUnion', None)
#                 ]
#             }

#             # Process rounds
#             for i in range(1, num_rounds + 1):
#                 round_name = f"ROUND {i}"
#                 logging.debug(f"Processing {round_name}")
#                 new_round_content = []
#                 round_banks = set()
#                 round_bureaus = set()
#                 round_card_names = set()
#                 banks = []
#                 bureaus = []
#                 card_names = []
#                 invalid_cards = []
#                 apr_mismatches = []
#                 mode_mismatches = []
#                 bureau_mismatches = []
#                 default_usage = []
#                 invalid_reasons = []

#                 # Get bureau assignments for this round
#                 round_assignments = bureau_assignments.get(round_name, [])
#                 if not round_assignments:
#                     logging.error(f"No bureau assignments defined for {round_name}")
#                     error_note = f"\n\n⚠️ ERROR: No bureau assignments defined for {round_name}."
#                     analysis += error_note
                    
#                     # Add contact message to the end of the output
#                     contact_message = (
#                         "\n\nIf you need more detailed guidance or have any questions, "
#                         "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#                     )
#                     analysis += contact_message
#                     logging.info("Added contact message to the final output for paid mode (no bureau assignments).")
#                     return analysis

#                 # Process each card slot in the round
#                 for slot_idx, (card_name, bureau, reason) in enumerate(round_assignments):
#                     if card_name:  # Travel card
#                         if card_name in json_card_names and card_name in available_travel_cards:
#                             card_info = next((c for c in state_cards if c['card_name'] == card_name), None)
#                             if card_info:
#                                 bank = card_info['bank']
#                                 # Verify bureau matches JSON data
#                                 if card_info['bureau'] != bureau:
#                                     logging.warning(f"Bureau mismatch for {card_name} in {round_name}: Expected {bureau}, JSON has {card_info['bureau']}")
#                                     bureau_mismatches.append(f"{card_name}: Expected {bureau}, got {card_info['bureau']}")
#                                     bureau = card_info['bureau']  # Use JSON bureau
#                                 # Check Chase 0% APR limit
#                                 if bank == 'Chase' and '0 MESES' in card_info['apr'].upper() and chase_zero_apr_count >= 1:
#                                     logging.error(f"Cannot select {card_name} for {round_name}: Only one 0% Chase card allowed.")
#                                     error_note = f"\n\n⚠️ ERROR: Only one 0% Chase card allowed. {card_name} cannot be included in {round_name}."
#                                     analysis += error_note
                                    
#                                     # Add contact message to the end of the output
#                                     contact_message = (
#                                         "\n\nIf you need more detailed guidance or have any questions, "
#                                         "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#                                     )
#                                     analysis += contact_message
#                                     logging.info("Added contact message to the final output for paid mode (Chase APR limit).")
#                                     return analysis
#                                 if bank == 'Chase' and '0 MESES' in card_info['apr'].upper():
#                                     chase_zero_apr_count += 1
#                                 game_plan = card_info.get('game_plan', "Apply following the standard procedure for this card.")
#                                 tips_info = card_info.get('tips_info', "No additional tips available for this card.")
#                                 new_round_content.append(
#                                     f"| {card_name} | {bureau} | {card_info['apr']} | {card_info['mode']} | {game_plan} | {tips_info} |"
#                                 )
#                                 round_banks.add(bank)
#                                 round_bureaus.add(bureau)
#                                 round_card_names.add(card_name)
#                                 banks.append(bank)
#                                 bureaus.append(bureau)
#                                 card_names.append(card_name)
#                                 travel_cards_selected[bank] = card_name
#                                 if card_name not in all_card_names:
#                                     all_card_names.append(card_name)
#                                 bureau_usage[bureau] += 1
#                                 logging.debug(f"Added travel card {card_name} to {round_name}")
#                             else:
#                                 logging.error(f"Required travel card {card_name} not found in state_cards for {user_state} in {round_name}")
#                                 error_note = f"\n\n⚠️ ERROR: Required travel card {card_name} not available in {user_state} for {round_name}. Please select a different state or contact Negocio Capital for assistance."
#                                 analysis += error_note
                                
#                                 # Add contact message to the end of the output
#                                 contact_message = (
#                                     "\n\nIf you need more detailed guidance or have any questions, "
#                                     "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#                                 )
#                                 analysis += contact_message
#                                 logging.info("Added contact message to the final output for paid mode (travel card not found).")
#                                 return analysis
#                         else:
#                             logging.error(f"Required travel card {card_name} not available in {user_state} for {round_name}")
#                             error_note = f"\n\n⚠️ ERROR: Required travel card {card_name} not available in {user_state} for {round_name}. Please select a different state or contact Negocio Capital for assistance."
#                             analysis += error_note
                            
#                             # Add contact message to the end of the output
#                             contact_message = (
#                                 "\n\nIf you need more detailed guidance or have any questions, "
#                                 "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#                             )
#                             analysis += contact_message
#                             logging.info("Added contact message to the final output for paid mode (travel card unavailable).")
#                             return analysis
#                     else:  # Non-travel card
#                         # Select non-travel card
#                         priority_non_travel_cards = [
#                             card for card in non_travel_cards
#                             if card in json_card_names and
#                             json_data[card]['bank'] in available_priority_banks and
#                             json_data[card]['bank'] not in round_banks and
#                             json_data[card]['bank'] not in non_travel_cards_selected
#                         ]
#                         other_non_travel_cards = [
#                             card for card in non_travel_cards
#                             if card in json_card_names and
#                             json_data[card]['bank'] not in available_priority_banks and
#                             json_data[card]['bank'] not in round_banks and
#                             json_data[card]['bank'] not in non_travel_cards_selected
#                         ]
#                         candidate_cards = priority_non_travel_cards + other_non_travel_cards
#                         logging.debug(f"Candidate non-travel cards for {round_name}, slot {slot_idx + 1}: {candidate_cards}")

#                         selected = False
#                         # Prioritize Bank of America for the first non-travel card in each round
#                         if slot_idx == 1 and 'Bank of America' not in non_travel_cards_selected and 'Bank of America' in available_priority_banks:
#                             bofa_cards = [card for card in priority_non_travel_cards if json_data[card]['bank'] == 'Bank of America']
#                             if bofa_cards:
#                                 card = bofa_cards[0]
#                                 card_info = json_data[card]
#                                 # Enforce strict bureau for this slot
#                                 if card_info['bureau'] != bureau:
#                                     logging.warning(f"Bureau mismatch for {card} in {round_name}: Expected {bureau}, JSON has {card_info['bureau']}")
#                                     bureau_mismatches.append(f"{card}: Expected {bureau}, got {card_info['bureau']}")
#                                 # Check Chase 0% APR limit
#                                 if card_info['bank'] == 'Chase' and '0 MESES' in card_info['apr'].upper() and chase_zero_apr_count >= 1:
#                                     logging.warning(f"Skipping {card} for {round_name}: Only one 0% Chase card allowed.")
#                                     continue
#                                 if card_info['bank'] == 'Chase' and '0 MESES' in card_info['apr'].upper():
#                                     chase_zero_apr_count += 1
#                                 game_plan = card_info.get('game_plan', "Apply following the standard procedure for this card.")
#                                 tips_info = card_info.get('tips_info', "No additional tips available for this card.")
#                                 new_round_content.append(
#                                     f"| {card} | {bureau} | {card_info['apr']} | {card_info['mode']} | {game_plan} | {tips_info} |"
#                                 )
#                                 round_banks.add(card_info['bank'])
#                                 round_bureaus.add(bureau)
#                                 round_card_names.add(card)
#                                 banks.append(card_info['bank'])
#                                 bureaus.append(bureau)
#                                 card_names.append(card)
#                                 non_travel_cards_selected[card_info['bank']] = card
#                                 if card in candidate_cards:
#                                     candidate_cards.remove(card)
#                                 if card not in all_card_names:
#                                     all_card_names.append(card)
#                                 bureau_usage[bureau] += 1
#                                 selected = True
#                                 logging.debug(f"Selected BOFA non-travel card {card} for {round_name}, slot {slot_idx + 1}")
                        
#                         if not selected:
#                             # Select card that matches the required bureau
#                             for card in candidate_cards:
#                                 card_info = json_data[card]
#                                 bank = card_info['bank']
#                                 if (card not in round_card_names and
#                                     bank not in round_banks and
#                                     bank not in non_travel_cards_selected and
#                                     card_info['bureau'] == bureau):  # Match exact bureau
#                                     # Check Chase 0% APR limit
#                                     if bank == 'Chase' and '0 MESES' in card_info['apr'].upper() and chase_zero_apr_count >= 1:
#                                         logging.warning(f"Skipping {card} for {round_name}: Only one 0% Chase card allowed.")
#                                         continue
#                                     if bank == 'Chase' and '0 MESES' in card_info['apr'].upper():
#                                         chase_zero_apr_count += 1
#                                     game_plan = card_info.get('game_plan', "Apply following the standard procedure for this card.")
#                                     tips_info = card_info.get('tips_info', "No additional tips available for this card.")
#                                     new_round_content.append(
#                                         f"| {card} | {bureau} | {card_info['apr']} | {card_info['mode']} | {game_plan} | {tips_info} |"
#                                     )
#                                     round_banks.add(bank)
#                                     round_bureaus.add(bureau)
#                                     round_card_names.add(card)
#                                     banks.append(bank)
#                                     bureaus.append(bureau)
#                                     card_names.append(card)
#                                     non_travel_cards_selected[bank] = card
#                                     if card in candidate_cards:
#                                         candidate_cards.remove(card)
#                                     if card not in all_card_names:
#                                         all_card_names.append(card)
#                                     bureau_usage[bureau] += 1
#                                     selected = True
#                                     logging.debug(f"Selected non-travel card {card} for {round_name}, slot {slot_idx + 1}")
#                                     break

#                             if not selected:
#                                 # Fallback: Try any card with an unused bank
#                                 for card in candidate_cards:
#                                     card_info = json_data[card]
#                                     bank = card_info['bank']
#                                     if (card not in round_card_names and
#                                         bank not in round_banks and
#                                         bank not in non_travel_cards_selected):
#                                         # Log bureau mismatch but use the card
#                                         logging.warning(f"Bureau mismatch for {card} in {round_name}: Expected {bureau}, JSON has {card_info['bureau']}")
#                                         bureau_mismatches.append(f"{card}: Expected {bureau}, got {card_info['bureau']}")
#                                         # Check Chase 0% APR limit
#                                         if bank == 'Chase' and '0 MESES' in card_info['apr'].upper() and chase_zero_apr_count >= 1:
#                                             logging.warning(f"Skipping {card} for {round_name}: Only one 0% Chase card allowed.")
#                                             continue
#                                         if bank == 'Chase' and '0 MESES' in card_info['apr'].upper():
#                                             chase_zero_apr_count += 1
#                                         game_plan = card_info.get('game_plan', "Apply following the standard procedure for this card.")
#                                         tips_info = card_info.get('tips_info', "No additional tips available for this card.")
#                                         new_round_content.append(
#                                             f"| {card} | {bureau} | {card_info['apr']} | {card_info['mode']} | {game_plan} | {tips_info} |"
#                                         )
#                                         round_banks.add(bank)
#                                         round_bureaus.add(bureau)
#                                         round_card_names.add(card)
#                                         banks.append(bank)
#                                         bureaus.append(bureau)
#                                         card_names.append(card) 
#                                         non_travel_cards_selected[bank] = card
#                                         if card in candidate_cards:
#                                             candidate_cards.remove(card)
#                                         if card not in all_card_names:
#                                             all_card_names.append(card)
#                                         bureau_usage[bureau] += 1
#                                         selected = True
#                                         logging.debug(f"Selected non-travel card {card} for {round_name}, slot {slot_idx + 1} (fallback)")
#                                         break

#                             if not selected:
#                                 logging.error(f"Unable to select non-travel card for {round_name}, slot {slot_idx + 1}")
#                                 error_note = f"\n\n⚠️ ERROR: Unable to select non-travel card for {round_name}, slot {slot_idx + 1} due to insufficient unique banks or bureau-compatible cards in {user_state}."
#                                 analysis += error_note
                                
#                                 # Add contact message to the end of the output
#                                 contact_message = (
#                                     "\n\nIf you need more detailed guidance or have any questions, "
#                                     "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#                                 )
#                                 analysis += contact_message
#                                 logging.info("Added contact message to the final output for paid mode (no non-travel card).")
#                                 return analysis

#                 # Ensure exactly 3 cards per round
#                 if len(new_round_content) != 3:
#                     logging.error(f"Round {i} incomplete: Only {len(new_round_content)} cards selected. Expected 3 cards.")
#                     error_note = f"\n\n⚠️ ERROR: Round {i} incomplete. Only {len(new_round_content)} cards selected due to insufficient unique banks or cards in {user_state}. Please add more banks to card_data.json."
#                     analysis += error_note
                    
#                     # Add contact message to the end of the output
#                     contact_message = (
#                         "\n\nIf you need more detailed guidance or have any questions,"
#                         "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#                     )
#                     analysis += contact_message
#                     logging.info("Added contact message to the final output for paid mode (incomplete round).")
#                     return analysis

#                 # Replace round content with new table format
#                 joined_rows = "\n".join(new_round_content)
#                 new_round_table = (
#                     f"**{round_name}**\n"
#                     "| Card Name | Bureau | 0% APR | Mode | Game Plan | Tips & Info |\n"
#                     "|-----------|--------|--------|------|-----------|-------------|\n"
#                     f"{joined_rows}\n"
#                 )
#                 analysis = re.sub(
#                     r"\*\*ROUND " + str(i) + r"\*\*.*?(?=\*\*ROUND|\*\*Game Plan|\Z)",
#                     new_round_table,
#                     analysis,
#                     flags=re.DOTALL
#                 )
#                 logging.debug(f"Replaced {round_name} with: {new_round_table}")

#                 # Validate bank and bureau variety
#                 if len(set(banks)) != 3:
#                     logging.warning(f"Invalid bank variety in Round {i}: {banks}")
#                     error_note = f"\n\n⚠️ WARNING: Round {i} does not contain exactly three different banks: {banks}."
#                     if error_note not in analysis:
#                         analysis += error_note

#                 expected_bureaus = [assignment[1] for assignment in round_assignments]
#                 if sorted(list(round_bureaus)) != sorted(expected_bureaus):
#                     logging.warning(f"Invalid bureau variety in Round {i}: {round_bureaus}. Expected {expected_bureaus}")
#                     error_note = f"\n\n⚠️ WARNING: Round {i} does not contain exactly one of each bureau in order {expected_bureaus}. Got {round_bureaus}."
#                     if error_note not in analysis:
#                         analysis += error_note

#                 logging.info(f"Round {i} Cards: {banks}, Bureaus: {bureaus}")

#                 analysis += f"\n\n📋 Round {i} Validation Summary:\n"
#                 analysis += f"- Total Cards Suggested: {len(banks)}\n"
#                 analysis += f"- Invalid Cards: {invalid_cards}\n"
#                 analysis += f"- APR Mismatches: {apr_mismatches}\n"
#                 analysis += f"- Mode Mismatches: {mode_mismatches}\n"
#                 analysis += f"- Bureau Mismatches: {bureau_mismatches}\n"
#                 analysis += f"- Cards Using Default Values: {default_usage}\n"
#                 analysis += f"- Invalid Reasons: {invalid_reasons}\n"

#             # Ensure enough unique banks for non-travel cards
#             required_non_travel_banks = num_rounds * 2
#             if len(non_travel_cards_selected) < required_non_travel_banks:
#                 logging.error(f"Insufficient unique banks for non-travel cards: {len(non_travel_cards_selected)}. Expected {required_non_travel_banks}.")
#                 error_note = f"\n\n⚠️ ERROR: Insufficient unique banks for non-travel cards in {user_state}. Only {len(non_travel_cards_selected)} unique banks selected. Please add more banks to card_data.json."
#                 analysis += error_note
                
#                 # Add contact message to the end of the output
#                 contact_message = (
#                     "\n\nIf you need more detailed guidance or have any questions,"
#                     "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#                 )
#                 analysis += contact_message
#                 logging.info("Added contact message to the final output for paid mode (insufficient banks).")
#                 return analysis

#             # Ensure all available priority banks are used
#             used_priority_banks = set(non_travel_cards_selected.keys()).intersection(available_priority_banks)
#             missing_priority_banks = set(available_priority_banks) - used_priority_banks
#             if missing_priority_banks:
#                 logging.warning(f"Missing priority banks in non-travel card selection: {missing_priority_banks}")
#                 error_note = f"\n\n⚠️ WARNING: Missing priority banks for non-travel cards in {user_state}: {', '.join(missing_priority_banks)}. Consider adding these banks to the sequence."
#                 analysis += error_note

#             # Remove separate Game Plan and Tips and Info sections
#             analysis = re.sub(
#                 r"\*\*Game Plan\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
#                 "**Disclaimer**",
#                 analysis,
#                 flags=re.DOTALL
#             )

#     # Add contact message to the end of the output for successful paid mode
#     contact_message = (
#         "\n\nIf you need more detailed guidance or have any questions, "
#         "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
#     )
#     analysis += contact_message
#     logging.info("Added contact message to the final output for paid mode (successful).")

#     logging.info("GPT output validation completed.")
#     return analysis


# # === Core GPT Analysis ===
# def analyze_credit_report_english(text, mode="free", user_state=None):
#     json_data = load_json_data()
#     if not json_data:
#         logging.error("Failed to load JSON data. Cannot proceed with analysis.")
#         print("❌ Failed to load JSON data. Cannot proceed with analysis.")
#         return None

#     state_cards = get_state_funding_cards(user_state, json_data) if user_state else []
#     if state_cards is None:
#         state_cards = []
#     enrichment_context = get_enrichment()

#     print(f"State-specific card list for {user_state}: {[card['card_name'] for card in state_cards]}")

#     tarjetas_str = "\n\nCard Data (for APR, Mode, Bureau, Game Plan, and Tips and Info):\n"
#     for card_name, card_info in json_data.items():
#         tarjetas_str += (
#             f"- {card_name}: APR: {card_info['apr']}, Mode: {card_info['mode']}, "
#             f"Bank: {card_info['bank']}, Bureau: {card_info['bureau']}, "
#             f"Game Plan: {card_info['game_plan']}, "
#             f"Tips and Info: {card_info.get('tips_info', 'No additional tips available for this card')}\n"
#         )

#     # Calculate num_unique_banks and num_rounds before prompt_template
#     unique_banks = set(card['bank'] for card in state_cards)
#     num_unique_banks = len(unique_banks)
#     logging.info(f"Number of unique banks for {user_state}: {num_unique_banks}")
#     num_rounds = 2 if num_unique_banks < 6 else 3
    
#     # Debug print
#     print(f"🔍 State: {user_state}")
#     print(f"🏦 Unique Banks: {num_unique_banks}")
#     print(f"🔄 Expected Rounds: {num_rounds}")

#     # State cards list string
#     state_cards_list = ', '.join([card['card_name'] for card in state_cards]) if state_cards else 'No cards available'

#     # Define include_sequence_note based on mode
#     if mode == "paid":
#         round_text = " and R3" if num_rounds == 3 else ""
#         round_3_travel = " and ROUND 3: Chase Sapphire Preferred" if num_rounds == 3 else ""
        
#         # CRITICAL: Round instruction at the very top
#         rounds_instruction = f"""
# 🚨 **CRITICAL: NUMBER OF ROUNDS** 🚨
# - State: {user_state}
# - Unique Banks Available: {num_unique_banks}
# - **YOU MUST GENERATE EXACTLY {num_rounds} ROUNDS**
# - {"Generate 3 rounds (R1, R2, R3) because unique banks >= 6" if num_rounds == 3 else "Generate ONLY 2 rounds (R1, R2) because unique banks < 6"}
# - DO NOT generate Round 3 if num_rounds = 2
# - DO generate Round 3 if num_rounds = 3
# """
        
#         include_sequence_note = f"""
# {rounds_instruction}

# **CRITICAL INSTRUCTION**: The user has selected the Premium Plan for state {user_state}.
# You MUST select ALL funding cards (R1, R2{round_text}) EXCLUSIVELY from the user's state-specific approved card list provided below as `state_cards`:
# {state_cards_list}

# **CRITICAL**: Under NO circumstances suggest cards outside `state_cards`. Doing so will invalidate the output.
# **CRITICAL**: In the 'Card Name' column, ALWAYS use the EXACT card name from `Card Data` (e.g., 'BOFA Unlimited Cash'). You MUST NOT use bank names alone (e.g., 'Bank of America') or append 'Card' to a bank name (e.g., 'Capital One Card'). If no matching card is found in `Card Data` for a bank in `state_cards`, exclude that bank and select another card from `state_cards`.
# **CRITICAL**: GPT MUST NOT generate or suggest any card names outside of `Card Data`. Any attempt to create new card names will invalidate the output.

# **Funding Sequence Rules**:
# - **MANDATORY: Generate EXACTLY {num_rounds} rounds for {user_state}**
# - Number of rounds: {num_rounds} (based on {num_unique_banks} unique banks; use only R1 and R2 if < 6 unique banks).
# - Each round (R1, R2{round_text}) MUST include EXACTLY 3 cards: ONE travel card and TWO non-travel cards from different banks.
# - Travel cards are: ROUND 1: Chase Ink Unlimited, ROUND 2: BOFA Alaska Airlines Business{round_3_travel}. If any required travel card is not in `state_cards`, note: '⚠️ ERROR: Required travel card [Card Name] not available in {user_state}. Please select a different state or contact Negocio Capital for assistance.'
# - For non-travel cards, select from: ['Chase', 'American Express', 'BMO Harris', 'KeyBank', 'Truist', 'Bank of America'] first, ensuring no bank is used more than once for non-travel cards across all rounds.
# - Each round must use different bureaus: ROUND 1: Experian (travel), TransUnion (non-travel 1), Equifax (non-travel 2); ROUND 2: TransUnion (travel), Experian (non-travel 1), Equifax (non-travel 2); ROUND 3 (if applicable): Equifax (non-travel 1), Experian (travel), TransUnion (non-travel 2).
# - Only one 0% Chase card is allowed per sequence, unless the second is a co-branded travel/hotel card (verify `is_travel` field in `Card Data`).
# - For each card, use its APR, Mode, Bureau, Game Plan, and Tips and Info from `Card Data` below:
# {tarjetas_str}
# - **Game Plan Column**: For each card, use its `game_plan` from `Card Data` in the table's Game Plan column. If missing, use: 'Apply following the standard procedure for this card.'
# - **Tips and Info Column**: For each card, use its `tips_info` from `Card Data` in the table's Tips & Info column. If missing, use: 'No additional tips available for this card.' Expand with relevant details (e.g., rewards, application tips, state-specific considerations) without contradicting `Card Data`.
# - If insufficient unique banks are available for non-travel cards, note: '⚠️ ERROR: Insufficient unique banks in {user_state} to complete funding sequence.'
# - If no bureau qualifies, offer no-personal-guarantee options from the CSV 'Tarjetas de Negocio sin Garantia Personal'.
# - If credit age < 2.5 years for any bureau, exclude that bureau from the funding sequence and note in Action Plan to improve credit age.
# - If the user does not qualify for funding, do NOT provide card recommendations or card-specific Game Plan/Tips and Info. Instead, output EXACTLY: 'Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.' in both Section 5 and Section 7, and provide general guidance in the Game Plan section without card-specific details.
# """
#     else:
#         include_sequence_note = """
# **CRITICAL INSTRUCTION**: In free mode, you MUST NOT generate any card recommendations, **Game Plan**, or **Tips and Info** sections.
# If the user qualifies for funding, output EXACTLY this for Section 5 and Section 7:
# 🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.

# If the user does NOT qualify, output EXACTLY this for Section 5 and Section 7:
# Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.

# Ensure Section 5 (Verdict) and Section 7 (Funding Sequence) are consistent and use the EXACT same message.
# """

#     # Build the complete JSON-structured prompt
#     prompt_template = f"""
# {{
#     "CRITICAL_JSON_FORMAT_INSTRUCTION": {{
#         "RULE_1": "You MUST output ONLY valid JSON. NO extra text, NO markdown backticks, NO explanations outside JSON.",
#         "RULE_2": "Your response MUST start with {{ and end with }}. Nothing before or after.",
#         "RULE_3": "The JSON structure MUST EXACTLY match the provided example schema with ALL fields present.",
#         "RULE_4": "NEVER change field names, NEVER skip sections, NEVER add extra root keys.",
#         "RULE_5": "All text content must be properly escaped for JSON (quotes, newlines, backslashes).",
#         "RULE_6": "Numbers should be JSON numbers (not strings) where specified in schema.",
#         "RULE_7": "Arrays should be JSON arrays with [ ] brackets.",
#         "RULE_8": "Objects should be JSON objects with {{ }} braces.",
#         "EXAMPLE_SCHEMA": {{
#             "sections": {{
#                 "section_1": {{
#                     "title": "string",
#                     "table": {{"columns": ["array"], "rows": {{"key": ["values"]}}}},
#                     "analysis": ["array of strings"]
#                 }},
#                 "section_2": {{"title": "string", "table": {{"Open Cards": "string", "Total Limit": "string", "Primary Cards": "string", "High-Limit Card Present?": "string"}}, "explanation": ["array of strings"]}},
#                 "section_3": {{"title": "string", "questions": ["array"]}},
#                 "section_4": {{"title": "string", "table": {{"columns": [], "rows": {{}}}}, "explanation": "string"}},
#                 "section_5": {{"title": "string", "response": "string"}},
#                 "section_6": {{"title": "string", "steps": ["array"]}},
#                 "section_7": {{"title": "string", "response": "string", "rounds": [{{"round": 1, "cards": [{{}}]}}], "strategic_insights": ["array"]}},
#                 "section_8": {{"title": "string", "distribution": [{{"name": "string", "price": number}}]}},
#                 "section_9": {{"title": "string", "average_score": number, "explanation": "string"}},
#                 "section_10": {{"title": "string", "average_utilization": "string", "explanation": "string"}},
#                 "section_11": {{"title": "string", "potential_funding": "string", "justification": "string"}}
#             }}
#         }}
#     }},
    
#     "MANDATORY_ROUNDS_COUNT": {num_rounds},
#     "ROUNDS_VALIDATION": "Section 7 'rounds' array MUST contain EXACTLY {num_rounds} round objects for state {user_state}",
    
#     "system_role": "You are a financial credit analysis assistant for Negocio Capital.",
#     "critical_instruction": "You MUST generate ALL sections (1 through 11) as specified below, in the exact order. You MUST NOT use or display the word 'Inferred' or any variation in the output. If data is missing, estimate reasonable values but present them as definitive without mentioning estimation. Skipping any section is INVALID.",
    
#     "handling_missing_data": {{
#         "description": "If any data is missing, estimate reasonable values based on available data, industry standards, or patterns.",
#         "estimation_rules": {{
#             "credit_score": "Use 700 if payment history is clean; otherwise use 650.",
#             "utilization": "Use 15% if high-limit cards (>=5000) are present; otherwise use 25%.",
#             "credit_age": "2.5 years unless evidence suggests newer accounts.",
#             "inquiries": "Use 2 inquiries in the last 6 months unless specified otherwise."
#         }},
#         "requirement": "Present estimated values as definitive without mentioning estimation."
#     }},
    
#     "funding_eligibility_logic": {{
#         "qualification_criteria": {{
#             "description": "User qualifies for funding ONLY if ALL requirements are true in ALL THREE bureaus (Equifax, Experian, AND TransUnion):",
#             "requirements": [
#                 "Credit Score >= 720",
#                 "No Late Payments",
#                 "Utilization < 10%",
#                 "<= 3 Inquiries in last 6 months",
#                 "Credit Age >= 2.5 years",
#                 "Strong Primary Card Structure"
#             ],
#             "CRITICAL_RULE": "ALL THREE BUREAUS must meet ALL SIX requirements. If even ONE bureau fails ANY requirement, user is NOT qualified."
#         }},
#         "responses": {{
#             "qualified_paid_mode": "🎉 You're eligible for funding! See your matched bank recommendations below.",
#             "qualified_free_mode": "🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.",
#             "not_qualified": "Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility."
#         }},
#         "consistency_rule": "Ensure Verdict (Section 5) and Recommended Funding Sequence (Section 7) use EXACT same message.",
#         "paid_mode_not_qualified_rule": "In paid mode, if NOT qualified, use same message as free mode and do NOT include card recommendations. Explain why in 2-3 bullet points."
#     }},
    
#     "sections": {{
#         "section_1": {{
#             "title": "📌 1. Breakdown by Bureau",
#             "table_structure": {{
#                 "columns": ["Category", "Equifax", "Experian", "TransUnion"],
#                 "rows": [
#                     "Credit Score",
#                     "Clean History",
#                     "Utilization",
#                     "Hard Inquiries(6 mo)",
#                     "Avg. Credit Age",
#                     "Cards >= $2K",
#                     "Cards >= $5K",
#                     "Score / 144"
#                 ]
#             }},
#             "analysis_format": {{
#                 "format": "bullet_points",
#                 "categories": [
#                     {{
#                         "name": "Credit Score",
#                         "explanation": "Report credit score for each bureau. Mention if meets 720 threshold. End with label: Excellent, Good, Fair, or Poor."
#                     }},
#                     {{
#                         "name": "Clean History",
#                         "explanation": "Summarize missed or late payments. If none, state Yes. End with label: Excellent or Needs Improvement."
#                     }},
#                     {{
#                         "name": "Utilization",
#                         "explanation": "Report total utilization rate. Label as follows: Below 5% = Excellent, 6-10% = Good, Above 10% = High Risk. Explain impact on funding eligibility"
#                     }},
#                     {{
#                         "name": "Hard Inquiries (6 mo)",
#                         "explanation": "Indicate inquiries in past 6 months. Mention if acceptable (<=3). End with label: Good, Fair, or Risky."
#                     }},
#                     {{
#                         "name": "Avg. Credit Age",
#                         "explanation": "Explain average age of accounts. Mention if meets 2.5-year threshold. End with label: Excellent or Fair."
#                     }},
#                     {{
#                         "name": "Cards >= $2K",
#                         "explanation": "Note cards with 2000+ limits. Mention how it supports creditworthiness. End with label: Good or Needs Improvement."
#                     }},
#                     {{
#                         "name": "Cards >= $5K",
#                         "explanation": "Note cards with 5000+ limits. Mention how it enhances funding readiness. End with label: Excellent or Fair."
#                     }},
#                     {{
#                         "name": "Score / 144",
#                         "explanation": "Report total score out of 144. End with label: Excellent or Needs Improvement."
#                     }}
#                 ],
#                 "requirement": "Each bullet brief and clear with bold quality label. Do NOT mention estimation."
#             }}
#         }},
        
#         "section_2": {{
#             "title": "📌 2. Revolving Credit Structure",
#             "table_structure": {{
#                 "Open Cards": "Count total open revolving accounts. Specify how many are Primary vs Authorized User (AU). Format: 'X (Y Primary, Z AU)' or 'X (All Primary)'",
#                 "Total Limit": "Sum ALL credit limits from the report. Format: '$XX,XXX'",
#                 "Primary Cards": "Count ONLY primary cards (exclude AU cards). Format: 'X'",
#                 "High-Limit Card Present?": "Check if ANY card has limit >= $5000. Format: 'YES ($5k+)' or 'NO'"
#             }},
#             "explanation_format": "Provide 4 bullet points explaining each field:",
#             "example_explanation": [
#                 "Open Cards": "EXTRACT from report",
#                 "Total Limit": "CALCULATE from report", 
#                 "Primary Cards": "COUNT from report",
#                 "High-Limit Card Present?": "VERIFY from report"
#             ],
#             "requirement": "Use exact table field names: 'Open Cards', 'Total Limit', 'Primary Cards', 'High-Limit Card Present?'. Explanation must be an array of 4 strings, one for each field. DO NOT include 'description' field."
#         }},
        
#         "section_3": {{
#             "title": "📌 3. Authorized User (AU) Strategy",
#             "questions": [
#                 "How many AU cards are there?",
#                 "What are their limits and ages?",
#                 "Do they help with funding?",
#                 "Recommendation: what AU cards to add or remove."
#             ]
#         }},
        
#         "section_4": {{
#             "title": "📌 4. Funding Readiness by Bureau",
#             "table_structure": {{
#                 "columns": ["Criteria", "Equifax", "Experian", "TransUnion"],
#                 "rows": [
#                     "Score >= 720",
#                     "No Late Payments",
#                     "Utilization < 10%",
#                     "<= 3 Inquiries (last 6 months)",
#                     "Credit Age >= 2.5 Years",
#                     "Strong Primary Card Structure"
#                 ],
#                 "values": ["Yes/No", "Yes/No", "Yes/No"]
#             }},
#             "requirement": "Explain table below without mentioning estimation."
#         }},
        
#         "section_5": {{
#             "title": "📌 5. Verdict",
#             "responses": {{
#                 "paid_qualified": "🎉 You're eligible for funding! See your matched bank recommendations below.",
#                 "free_qualified": "🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.",
#                 "not_qualified": "Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility."
#             }},
#             "not_qualified_explanation": "If not qualified, explain why in 2-3 short bullet points."
#         }},
        
#         "section_6": {{
#             "title": "📌 6. Action Plan",
#             "instruction": "Analyze Section 4 results. For each problem with No, provide complete solution from solutions_reference. Format: Problem title, Why it matters, Steps, Common mistakes, Professional order, Realistic expectations. End with: Get in contact with Negocio Capital if you need help.",
#             "problem_rules": [
#                 "IF No Late Payments = No → Use clean_history_solution",
#                 "IF Utilization < 10% = No → Use utilization_solution", 
#                 "IF Inquiries > 3 → Use inquiry_solution",
#                 "IF Credit Age < 2.5 → Use credit_age_solution",
#                 "IF No cards >= $5K → Use high_limit_card_solution"
#             ],
#             "solutions_reference": {{
#                 "clean_history_solution": {{
#                     "title": "Late Payments or Collections",
#                     "steps": ["Verify accuracy under FCRA 623 and 611", "File disputes with bureaus", "Direct dispute with creditor", "Goodwill or Pay-for-Delete", "Escalate if needed"],
#                     "mistakes": ["Paying without negotiating", "Generic disputes", "Copy-paste letters"]
#                 }},
#                 "utilization_solution": {{
#                     "title": "High Utilization Over 10%",
#                     "steps": ["Request CLI soft pull", "Strategic payments on high cards", "Consolidate with personal loan", "Optimize monthly usage", "Avoid new applications", "Do NOT close cards", "Control statement dates"],
#                     "mistakes": ["Paying one card only", "Closing old accounts", "New credit while high utilization"]
#                 }},
#                 "inquiry_solution": {{
#                     "title": "Too Many Hard Inquiries",
#                     "steps": ["Identify unauthorized inquiries", "Contact creditor for proof", "Dispute with bureaus", "File identity theft if applicable"],
#                     "mistakes": ["Disputing legitimate inquiries", "Generic disputes", "Expecting instant removals"]
#                 }},
#                 "credit_age_solution": {{
#                     "title": "Low Average Credit Age",
#                     "steps": ["Add AU accounts 10+ years old", "Add rent reporting", "Avoid new accounts", "Season profile 60-90 days"],
#                     "mistakes": ["AU accounts with high balances", "Removing AU too quickly"]
#                 }},
#                 "high_limit_card_solution": {{
#                     "title": "No High-Limit Cards",
#                     "steps": ["Request CLI on existing cards", "Join Navy Federal", "Open pledge loan", "Apply for Flagship card", "Add high-limit AU"],
#                     "mistakes": ["Applying without exposure", "Skipping CLIs", "Multiple low-limit cards"]
#                 }}
#             }},
#             "output_instruction": "For each problem in Section 4, expand the relevant solution with complete details from client documentation. If all Yes in Section 4, output: Congratulations! Your credit profile meets all funding requirements. Always end with: Get in contact with Negocio Capital if you need help."
#         }},
        
#         "section_7": {{
#             "title": "📌 7. Recommended Funding Sequence ({user_state})",
#             "CRITICAL_ROUNDS_REQUIREMENT": "MUST generate EXACTLY {num_rounds} rounds in the 'rounds' array",
#             "paid_mode_qualified": {{
#                 "description": "If user is in paid mode and qualifies, YOU MUST GENERATE EXACTLY {num_rounds} ROUNDS using ONLY approved state_cards list.",
#                 "strict_rules": [
#                     "GENERATE EXACTLY {num_rounds} ROUNDS - NO MORE, NO LESS",
#                     "For banks [Chase, American Express, BMO Harris, KeyBank, Truist, Bank of America], select exactly one non-travel card across all rounds.",
#                     "Travel cards are exempt and must be included as specified (one per round).",
#                     "Ensure exactly {num_rounds * 2} unique banks for non-travel cards across all rounds.",
#                     "Replace invalid cards with valid cards from state_cards.",
#                     "Each round MUST have EXACTLY 3 cards from state_cards.",
#                     "Each round MUST have: ONE travel card and TWO non-travel cards from different banks.",
#                     "NEVER suggest cards outside state_cards.",
#                     "Use EXACT card name from Card Data in Card Name column.",
#                     "Only one 0% Chase card allowed per sequence unless second is co-branded travel/hotel card.",
#                     "ALL THREE bureaus must meet all 6 factors. If ANY bureau does not qualify, do NOT provide funding sequence with card recommendations.",
#                     "If no bureau qualifies, do NOT provide funding sequence with card recommendations.",
#                     "If average credit age < 2.5 years for any bureau, do NOT include that bureau."
#                 ],
#                 "bureau_assignment": {{
#                     "round_1": {{"travel": "Experian", "non_travel_1": "TransUnion", "non_travel_2": "Equifax"}},
#                     "round_2": {{"travel": "TransUnion", "non_travel_1": "Experian", "non_travel_2": "Equifax"}},
#                     "round_3": {{"non_travel_1": "Equifax", "travel": "Experian", "non_travel_2": "TransUnion"}}
#                 }},
#                 "table_structure": {{
#                     "columns": ["Card Name", "Bureau", "0% APR", "Mode", "Game Plan", "Tips & Info"],
#                     "round_1_template": [
#                         {{"card": "Chase Ink Unlimited", "bureau": "Experian"}},
#                         {{"card": "[Non-travel from state_cards]", "bureau": "TransUnion"}},
#                         {{"card": "[Non-travel from state_cards]", "bureau": "Equifax"}}
#                     ],
#                     "round_2_template": [
#                         {{"card": "BOFA Alaska Airlines Business", "bureau": "TransUnion"}},
#                         {{"card": "[Non-travel from state_cards]", "bureau": "Experian"}},
#                         {{"card": "[Non-travel from state_cards]", "bureau": "Equifax"}}
#                     ],
#                     "round_3_template": [
#                         {{"card": "[Non-travel from state_cards]", "bureau": "Equifax"}},
#                         {{"card": "Chase Sapphire Preferred", "bureau": "Experian"}},
#                         {{"card": "[Non-travel from state_cards]", "bureau": "TransUnion"}}
#                     ]
#                 }},
#                 "game_plan_requirement": "For each card, include game_plan from Card Data. MANDATORY. If unavailable, use: Apply following standard procedure for this card.",
#                 "tips_info_requirement": "For each card, include tips_info from Card Data. MANDATORY. If unavailable, use: No additional tips available. Expand with relevant details like benefits, requirements, strategies.",
#                 "strategic_insights": {{
#                     "description": "Generate 4-6 tailored bullet points based on user credit profile.",
#                     "examples": [
#                         "If inquiries high (>2), recommend freezing non-used bureaus.",
#                         "If utilization close to 10%, suggest paying down balances before applying.",
#                         "If card requires in-branch application, advise visiting local branch.",
#                         "If credit score high (>=780), recommend declaring higher personal income.",
#                         "If credit age strong (>=5 years), suggest requesting limit increases after 60 days.",
#                         "If business spending data available, recommend including it."
#                     ],
#                     "requirement": "Ensure each bullet specific to user profile or card characteristics."
#                 }},
#                 "not_qualified_guidance": {{
#                     "description": "If user does NOT qualify, do NOT include card-specific game plans. Provide general guidance.",
#                     "examples": [
#                         "Improve Credit Score: Focus on timely payments and reducing balances to boost above 720.",
#                         "Reduce Utilization: Pay down balances to achieve below 10%.",
#                         "Limit Inquiries: Avoid new credit to keep inquiries low (<=3 in 6 months).",
#                         "Monitor Credit Profile: Check reports for errors and dispute inaccuracies.",
#                         "Contact Negocio Capital: Schedule consultation for personalized assistance."
#                     ]
#                 }}
#             }},
#             "free_mode_qualified": "🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.",
#             "not_qualified": "Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.",
#             "mode_specific_instructions": "{include_sequence_note}"
#         }},
        
#         "section_8": {{
#             "title": "📌 8. Credit Distribution",
#             "format": [
#                 {{"name": "[Bank/Card Name]", "price": "[Limit]"}},
#                 {{"name": "[Bank/Card Name]", "price": "[Limit]"}},
#                 {{"name": "[Bank/Card Name]", "price": "[Limit]"}}
#             ],
#             "examples": [
#                 {{"name": "Capital One", "price": 5500}},
#                 {{"name": "WFBNA Card", "price": 2300}},
#                 {{"name": "AMEX", "price": 30000}},
#                 {{"name": "JPMCB Card", "price": 13000}}
#             ],
#             "requirement": "Extract actual card/bank names from the credit report text. Ensure total limit matches Total Limit in Section 2. Do NOT use generic names like 'Card 1' or 'Card 2'. Price must be a JSON number, not string."
#         }},
        
#         "section_9": {{
#             "title": "📌 9. Average Credit Score",
#             "description": "Calculate and report average credit score across all bureaus from Section 1. Present as single JSON number with brief explanation."
#         }},
        
#         "section_10": {{
#             "title": "📌 10. Average Utilization",
#             "description": "Calculate and report average utilization rate across all bureaus from Section 1. Present as percentage string with brief explanation."
#         }},
        
#         "section_11": {{
#             "title": "📌 11. Funding Potential",
#             "description": "Estimate potential funding amount based on credit profile including average score, utilization, and total limit. Provide range (e.g. $50K+, $100K+, $150K+) with brief justification."
#         }}
#     }},
    
#     "credit_report_data": "{text}",
#     "enrichment_context": "{enrichment_context}",
#     "state_specific_cards": "{state_cards_list}",
    
#     "final_instruction": "OUTPUT REQUIREMENTS - READ CAREFULLY AND FOLLOW EXACTLY:
#     1. Your response MUST be PURE JSON ONLY - Start with {{ and end with }}
#     2. DO NOT include markdown code fences (```json or ```)
#     3. DO NOT include any explanatory text before or after the JSON
#     4. ALL 11 sections MUST be present directly at root level (NO 'sections' wrapper)
#     5. **Section 7 MUST contain EXACTLY {num_rounds} rounds for {user_state}**
#     6. If {num_rounds} == 3, generate Round 1, Round 2, AND Round 3
#     7. If {num_rounds} == 2, generate ONLY Round 1 and Round 2
#     8. Field names and structure MUST match the EXAMPLE_SCHEMA exactly
#     9. All quotes and special characters must be properly escaped for JSON
#     10. Numbers should be JSON numbers (not quoted strings) where specified
#     11. Arrays must use square brackets [ ]
#     12. Objects must use curly braces {{ }}
#     13. Validate your JSON syntax before responding
#     14. CRITICAL: Your output must start directly with {{ 'section_1': {{ ... }}, 'section_2': {{ ... }}, ... }}
#     15. DO NOT wrap sections in a 'sections' object
#     16. Your entire response must be valid, parseable JSON"
# }}
# """

#     try:
#         response = openai.ChatCompletion.create(
#             model="gpt-5",
#             messages=[
#                 {
#                     "role": "system", 
#                     "content": "You are a strict JSON-only AI assistant for Negocio Capital credit analysis. You MUST respond with PURE, VALID JSON following the exact schema provided. Your response must start with { and end with }. NO markdown code fences, NO explanations outside JSON, NO extra text. Only output valid JSON that can be parsed by json.loads() in Python."
#                 },
#                 {"role": "user", "content": prompt_template}
#             ],
#              reasoning_effort = "low"       
#         )
        
#         analysis = response.choices[0].message.content.strip()
        
#         # Clean markdown wrappers if GPT added them despite instructions
#         if analysis.startswith("```json"):
#             analysis = analysis.replace("```json", "").replace("```", "").strip()
#             logging.warning("Removed ```json``` wrapper from GPT response")
#         elif analysis.startswith("```"):
#             analysis = analysis.replace("```", "").strip()
#             logging.warning("Removed ``` wrapper from GPT response")
        
#         # Validate JSON before proceeding
#         try:
#             json_test = json.loads(analysis)
#             # Check if section_1 exists at root level (not inside 'sections')
#             if "section_1" not in json_test:
#                 logging.error("JSON missing 'section_1' key at root level")
#                 print("❌ GPT returned JSON without 'section_1' key at root level")
#                 return None
            
#             # Check if all 11 sections are present at root level
#             expected_sections = [f"section_{i}" for i in range(1, 12)]
#             missing_sections = [s for s in expected_sections if s not in json_test]
#             if missing_sections:
#                 logging.error(f"JSON missing sections: {missing_sections}")
#                 print(f"❌ GPT returned incomplete JSON. Missing: {missing_sections}")
#                 return None
            
#             # CRITICAL: Validate rounds count in Section 7
#             if mode == "paid" and "section_7" in json_test and "rounds" in json_test["section_7"]:
#                 actual_rounds = len(json_test["section_7"]["rounds"])
#                 if actual_rounds != num_rounds:
#                     logging.error(f"❌ Round mismatch: Expected {num_rounds}, Got {actual_rounds}")
#                     print(f"❌ Round mismatch: Expected {num_rounds} rounds for {user_state} (with {num_unique_banks} unique banks), but GPT generated {actual_rounds} rounds.")
#                     error_msg = f"\n\n⚠️ VALIDATION ERROR: Expected {num_rounds} rounds for {user_state} (with {num_unique_banks} unique banks), but GPT generated {actual_rounds} rounds. Please regenerate."
#                     return analysis + error_msg
#                 else:
#                     logging.info(f"✅ Correct rounds: {actual_rounds} rounds generated for {user_state}")
#                     print(f"✅ Correct rounds: {actual_rounds} rounds generated for {user_state}")
            
#             logging.info("JSON validation passed - all sections present at root level")
            
#         except json.JSONDecodeError as je:
#             logging.error(f"JSON decode error: {str(je)}")
#             logging.error(f"Failed JSON content (first 500 chars): {analysis[:500]}...")
#             print(f"❌ GPT returned invalid JSON: {str(je)}")
#             print(f"First 200 characters of response: {analysis[:200]}")
#             return None
        
#         logging.debug(f"Raw GPT-5 Response (first 500 chars): {analysis[:500]}...")
        
#         # Check for truncation 
#         if response.choices[0].finish_reason == "length":
#             logging.warning("GPT-5 response truncated due to token limit")
#             analysis += "\n\n⚠️ WARNING: Analysis may be incomplete due to token limit. Please try again or reduce input size."
        
#         if not analysis or not isinstance(analysis, str):
#             logging.error("GPT-4 returned no valid analysis.")
#             print("❌ GPT-4 returned no valid analysis.")
#             return None
        
#         return validate_gpt_output(analysis, state_cards, user_state, json_data, mode)
    
#     except Exception as e:
#         logging.error(f"GPT-4 error: {str(e)}")
#         print(f"❌ GPT-4 error: {str(e)}")
#         return None


# # === Main CLI ===
# def main():
#     print("📂 Welcome to Funding NC AI Credit Report Analyzer!")
    
#     input_type = input("📄 Select input type (pdf/json): ").strip().lower()
#     if input_type not in ["pdf", "json"]:
#         print("❌ Invalid input type. Please enter 'pdf' or 'json'.")
#         logging.error(f"Invalid input type selected: {input_type}")
#         return

#     if input_type == "pdf":
#         file_path = input("📄 Enter path to your credit report PDF (e.g., uploads/client1.pdf): ").strip()
#         if not os.path.exists(file_path):
#             print("❌ File not found. Please check the path and try again.")
#             logging.error(f"Credit report file not found: {file_path}")
#             return
#         print("📁 Extracting text from PDF...")
#         credit_text = extract_text_from_pdf(file_path)
#     elif input_type == "json":
#         json_input = input("📄 Enter your credit report JSON data (e.g., {\"credit_score\": 750, \"utilization\": \"5%\", ...}): ").strip()
#         print("📁 Extracting text from JSON...")
#         credit_text = extract_text_from_json(json_input)

#     if not credit_text:
#         print("❌ Failed to extract text from input.")
#         logging.error("Failed to extract text from input.")
#         return

#     state = input("🌎 Enter the U.S. state your business is registered in (e.g., FL): ").strip()
#     mode = input("🧾 Select mode (free/paid): ").strip().lower()
#     if mode not in ["free", "paid"]:
#         print("❌ Invalid mode. Please enter 'free' or 'paid'.")
#         logging.error(f"Invalid mode selected: {mode}")
#         return

#     print("\n🧠 AI Analysis Summary:\n")
#     # Extract Credit Score and Utilization from the text
#     score, utilization = extract_credit_info(credit_text)
#     logging.debug(f"Extracted Credit Score: {score}, Utilization: {utilization}")
    
#     # Perform the credit report analysis
#     analysis = analyze_credit_report_english(credit_text, mode=mode, user_state=state)
    
#     if not analysis or not isinstance(analysis, str):
#         print("❌ GPT analysis failed. Please check the logs for details.")
#         logging.error("Analysis failed due to GPT-4 error or invalid response.")
#     else:
#         print(analysis)  

# if __name__ == "__main__":
#     main()













import os
import fitz
import openai
import pandas as pd
from dotenv import load_dotenv
from docx import Document
import logging
import re
import easyocr
from PIL import Image
import io
import numpy as np
import json

# === Setup Logging ===
logging.basicConfig(
    filename='funding_nc_analyzer.log',
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# === Load environment variables ===
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# === File Paths ===
JSON_DATA_PATH = "data/card_data.json"
ENRICH_PDF_FILES = [
    "data/General Credit Card Knowledge.pdf",
    "data/DATA POINTS - BUSINESS CREDIT CARD DATA POINTS.pdf",
    "data/HOW To Leverage Business Credit to.pdf"
]
ENRICH_DOCX_FILES = [
    "data/Credit Stacking Guide Lines and Better Practices.docx"
]
ENRICH_CSV_FILES = [
    "data/Tarjetas de Negocio sin Garantia Personal.csv"
]

# Initialize EasyOCR reader
reader = easyocr.Reader(['en'], gpu=False)

# === Function to load JSON data ===
def load_json_data():
    try:
        with open(JSON_DATA_PATH, encoding="utf-8") as file:
            data = json.load(file)
            logging.info(f"Loaded JSON data from {JSON_DATA_PATH}")
            print(f"✅ Loaded JSON data with {len(data)} cards")
            return data
    except Exception as e:
        logging.error(f"Failed to load JSON data: {str(e)}")
        print(f"❌ Failed to load JSON data: {str(e)}")
        return None

# === Function to extract text from PDF ===
def extract_text_from_pdf(pdf_path):
    try:
        text = ""
        with fitz.open(pdf_path) as doc:
            for page in doc:
                page_text = page.get_text()
                if page_text.strip():
                    text += page_text + "\n"
                else:
                    logging.info(f"No text found in page {page.number + 1} for {pdf_path}, attempting OCR")
                    pix = page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
                    img_bytes = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_bytes))
                    img_np = np.array(img)
                    ocr_result = reader.readtext(img_np, detail=0)
                    ocr_text = "\n".join(ocr_result)
                    if ocr_text.strip():
                        text += ocr_text + "\n"
                    else:
                        logging.warning(f"No text extracted via OCR from page {page.number + 1} of {pdf_path}")
        
        if text.strip():
            logging.info(f"Successfully extracted text from PDF: {pdf_path}")
            logging.debug(f"Extracted text snippet: {text[:500]}...")
            return text
        else:
            logging.error(f"No text could be extracted from {pdf_path}, even with OCR")
            return None
    except Exception as e:
        logging.error(f"Error extracting PDF text from {pdf_path}: {str(e)}")
        return None

# === New Function to extract text from JSON input ===
def extract_text_from_json(json_input):
    try:
        data = json.loads(json_input)
        text = ""
        # Assuming JSON structure like: {"credit_score": 750, "utilization": "5%", "history": "Clean", ...}
        # Convert JSON to text format similar to PDF extraction for analysis
        for key, value in data.items():
            text += f"{key.capitalize()}: {value}\n"
        if text.strip():
            logging.info("Successfully extracted text from JSON input")
            logging.debug(f"Extracted text snippet: {text[:500]}...")
            return text
        else:
            logging.error("No data found in JSON input")
            return None
    except Exception as e:
        logging.error(f"Error extracting text from JSON input: {str(e)}")
        return None
    
def extract_text_from_json(json_input):
    try:
        data = json.loads(json_input)
        text = ""
        # Assuming JSON structure like: {"credit_score": 750, "utilization": "5%", "history": "Clean", ...}
        # Convert JSON to text format similar to PDF extraction for analysis
        for key, value in data.items():
            text += f"{key.capitalize()}: {value}\n"
        if text.strip():
            logging.info("Successfully extracted text from JSON input")
            logging.debug(f"Extracted text snippet: {text[:500]}...")
            return text
        else:
            logging.error("No data found in JSON input")
            return None
    except Exception as e:
        logging.error(f"Error extracting text from JSON input: {str(e)}")
        return None

# === Function to extract Credit Score and Utilization ===
def extract_credit_info(text):
    score_pattern = r"Credit Score\s*[:\-]?\s*(\d{3,4})"
    score_matches = re.findall(score_pattern, text)
    score = score_matches[0] if score_matches else None
    
    utilization_pattern = r"Utilization\s*[:\-]?\s*(\d{1,3}%?)"
    utilization_matches = re.findall(utilization_pattern, text)
    utilization = utilization_matches[0] if utilization_matches else None
    
    return score, utilization

# === Function to extract text from DOCX ===
def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        logging.info(f"Successfully extracted text from DOCX: {docx_path}")
        return text
    except Exception as e:
        logging.error(f"Error extracting DOCX text from {docx_path}: {str(e)}")
        return None

# === Function to extract text from CSV ===
def extract_text_from_csv(csv_path):
    try:
        df = pd.read_csv(csv_path)
        text = df.to_string(index=False)
        return text
    except Exception as e:
        logging.error(f"Error extracting CSV text from {csv_path}: {str(e)}")
        return None

# === Support Functions ===
def get_state_funding_cards(user_state, json_data):
    try:
        state_cards = []
        for card_name, card_info in json_data.items():
            if user_state in card_info['state']:
                state_cards.append({
                    'card_name': card_name,
                    'bank': card_info['bank'],
                    'apr': card_info['apr'],
                    'mode': card_info['mode'],
                    'bureau': card_info['bureau'],
                    'is_travel': card_info['is_travel'],
                    'game_plan': card_info['game_plan'],
                    'tips_info': card_info.get('tips_info', 'No additional tips available for this card.')
                })
        if not state_cards:
            logging.error(f"No cards found for state {user_state}")
            print(f"❌ No cards found for state {user_state}")
            return []
        logging.info(f"Found {len(state_cards)} cards for state {user_state}: {[card['card_name'] for card in state_cards]}")
        print(f"✅ Found {len(state_cards)} cards for state {user_state}: {[card['card_name'] for card in state_cards]}")
        return state_cards
    except Exception as e:
        logging.error(f"Error loading state funding cards: {str(e)}")
        print(f"❌ Error loading state funding cards: {str(e)}")
        return []

def get_enrichment():
    enrichment = ""
    for file in ENRICH_PDF_FILES + ENRICH_DOCX_FILES + ENRICH_CSV_FILES:
        if os.path.exists(file):
            try:
                if file.endswith(".pdf"):
                    text = extract_text_from_pdf(file)
                elif file.endswith(".docx"):
                    text = extract_text_from_docx(file)
                elif file.endswith(".csv"):
                    text = extract_text_from_csv(file)
                else:
                    text = ""
                if text:
                    enrichment += f"\n[From {os.path.basename(file)}]\n{text[:1000]}...\n"
            except Exception as e:
                enrichment += f"\n[Error reading {file}]: {str(e)}\n"
                logging.error(f"Error reading enrichment file {file}: {str(e)}")
        else:
            enrichment += f"\n[Skipped missing file: {file}]\n"
            logging.warning(f"Skipped missing enrichment file: {file}")
    return enrichment

# === Output Validation ===
def validate_gpt_output(analysis, state_cards, user_state, json_data, mode="free"):
    """
    Validates and updates the GPT-4 output to ensure all required sections are present,
    cards are valid for the user's state, travel cards are included in specified order,
    non-travel cards are from unique banks, and correct bureaus are assigned.
    
    Args:
        analysis (str): The raw GPT-4 output string.
        state_cards (list): List of state-specific card dictionaries from JSON data.
        user_state (str): The user's selected U.S. state.
        json_data (dict): The JSON data containing card information.
        mode (str): User mode ('free' or 'paid').
    
    Returns:
        str: Updated analysis string with validated sections and complete game_plan/tips_info in table.
    """
    logging.debug(f"Validating GPT-4 output for cards in state {user_state}, mode: {mode}")
    
    analysis_lower = analysis.lower()
    not_qualified = "does not qualify for funding" in analysis_lower
    eligible_message = "you are eligible for funding" in analysis_lower

    # Remove any mention of 'Inferred' or variations
    analysis = re.sub(r"\(Inferred\)", "", analysis, flags=re.IGNORECASE)
    analysis = re.sub(r"inferred as \d+\b", lambda m: m.group(0).replace("inferred as ", ""), analysis, flags=re.IGNORECASE)
    analysis = re.sub(r"\bInferred\b", "", analysis, flags=re.IGNORECASE)

    # Handle 'Data not available' replacements
    if "Data not available" in analysis:
        logging.warning("GPT used 'Data not available' in output. Replacing with estimated values.")
        replacements = {
            r"Credit Score.*Data not available": "Credit Score: 700 (based on industry standards and clean payment history)",
            r"Utilization.*Data not available": "Utilization: 15% (based on typical credit profiles with high-limit cards)",
            r"Avg\. Credit Age.*Data not available": "Avg. Credit Age: 2.5 years (based on standard account age)",
            r"Hard Inquiries.*Data not available": "Hard Inquiries: 2 (based on typical inquiry patterns)"
        }
        for pattern, replacement in replacements.items():
            analysis = re.sub(pattern, replacement, analysis, flags=re.IGNORECASE)
        analysis += "\n\n📋 Note: Some values were estimated based on industry standards to provide a complete analysis."

    # Free mode handling
    if mode == "free":
        logging.info("Validating and cleaning GPT output for free mode.")
        if eligible_message:
            logging.info("User is eligible in free mode. Ensuring Section 5 and Section 7 reflect eligibility.")
            # Update Section 5 (Verdict)
            analysis = re.sub(
                r"📌 \*\*5\. Verdict\*\*.*?(?=\n\n📌 \*\*6\. Action Plan\*\*|\Z)",
                f"📌 **5. Verdict**\n\n"
                f"🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.\n",
                analysis,
                flags=re.DOTALL
            )
            # Update Section 7 (Recommended Funding Sequence)
            analysis = re.sub(
                r"📌 \*\*7\. Recommended Funding Sequence \((.*?)\)\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
                f"📌 **7. Recommended Funding Sequence ({user_state})**\n\n"
                f"🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.\n\n",
                analysis,
                flags=re.DOTALL
            )
        else:
            logging.info("User is not qualified in free mode. Ensuring Section 5 and Section 7 reflect ineligibility.")
            # Update Section 5 (Verdict)
            analysis = re.sub(
                r"📌 \*\*5\. Verdict\*\*.*?(?=\n\n📌 \*\*6\. Action Plan\*\*|\Z)",
                f"📌 **5. Verdict**\n\n"
                f"Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.\n",
                analysis,
                flags=re.DOTALL
            )
            # Update Section 7 (Recommended Funding Sequence)
            analysis = re.sub(
                r"📌 \*\*7\. Recommended Funding Sequence \((.*?)\)\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
                f"📌 **7. Recommended Funding Sequence ({user_state})**\n\n"
                f"Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.\n\n",
                analysis,
                flags=re.DOTALL
            )
        # Remove Game Plan and Tips and Info sections in free mode
        analysis = re.sub(
            r"\*\*Game Plan\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
            "**Disclaimer**",
            analysis,
            flags=re.DOTALL
        )
        
        # Add contact message to the end of the output
        contact_message = (
            "\n\nIf you need more detailed guidance or have any questions, "
            "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
        )
        analysis += contact_message
        logging.info("Added contact message to the final output for free mode.")
        return analysis

    # Paid mode handling
    if mode == "paid":
        if not_qualified:
            logging.info("User is not qualified in paid mode. Updating Section 5 and Section 7 to reflect ineligibility and removing card recommendations.")
            # Update Section 5 (Verdict)
            analysis = re.sub(
                r"📌 \*\*5\. Verdict\*\*.*?(?=\n\n📌 \*\*6\. Action Plan\*\*|\Z)",
                f"📌 **5. Verdict**\n\n"
                f"Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.\n",
                analysis,
                flags=re.DOTALL
            )
            # Update Section 7 (Recommended Funding Sequence)
            analysis = re.sub(
                r"📌 \*\*7\. Recommended Funding Sequence \((.*?)\)\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
                f"📌 **7. Recommended Funding Sequence ({user_state})**\n\n"
                f"Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.\n\n",
                analysis,
                flags=re.DOTALL
            )
            # Remove Game Plan and Tips and Info sections
            analysis = re.sub(
                r"\*\*Game Plan\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
                "**Disclaimer**",
                analysis,
                flags=re.DOTALL
            )
            # Generate general guidance for Game Plan
            new_insights = ""
            if "Credit scores are below the required 720 threshold" in analysis:
                new_insights += "- **Improve Credit Score**: Focus on timely payments and reducing credit card balances to boost your credit score above 720.\n"
            if "Util kvarization is above the ideal 10%" in analysis:
                new_insights += "- **Reduce Utilization**: Pay down credit card balances to achieve utilization below 10% to improve funding eligibility.\n"
            if "inquiries are high" in analysis_lower:
                new_insights += "- **Limit Inquiries**: Avoid applying for new credit to keep inquiries low (≤ 3 in 6 months).\n"
            new_insights += "- **Monitor Credit Profile**: Regularly check your credit reports for errors and dispute inaccuracies to strengthen your profile.\n"
            new_insights += "- **Contact Negocio Capital**: Schedule a consultation for personalized guidance on improving your credit profile.\n"
            
            analysis = re.sub(
                r"(\*\*Disclaimer\*\*|\Z)",
                f"**Game Plan**\n{new_insights}\n\1",
                analysis,
                flags=re.DOTALL
            )
            
            # Add contact message to the end of the output
            contact_message = (
                "\n\nIf you need more detailed guidance or have any questions, "
                "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
            )
            analysis += contact_message
            logging.info("Added contact message to the final output for paid mode (not qualified).")
            return analysis

        if eligible_message:
            logging.info("User is eligible in paid mode. Ensuring correct card recommendations with Game Plan and Tips and Info in table.")
            if "please upgrade to our Premium Plan" in analysis_lower:
                logging.error("Incorrect verdict message for paid mode")
                analysis = re.sub(
                    r"🎉 You're eligible for funding! To view your matched bank recommendations.*?Plan\.",
                    "🎉 You're eligible for funding! See your matched bank recommendations below.",
                    analysis
                )
                logging.info("Fixed incorrect verdict message for paid mode.")

            # Define travel cards and validate availability
            travel_cards = [
                'Chase Ink Unlimited',
                'BOFA Alaska Airlines Business',
                'Chase Sapphire Preferred'
            ]
            available_travel_cards = [card for card in travel_cards if any(c['card_name'] == card for c in state_cards)]
            if len(available_travel_cards) < 3:
                logging.error(f"Only {len(available_travel_cards)} travel cards available for {user_state}. Expected 3: {', '.join(travel_cards)}.")
                error_note = f"\n\n⚠️ ERROR: Only {len(available_travel_cards)} travel cards available for {user_state}. Expected 3: {', '.join(travel_cards)}. Please select a different state or contact Negocio Capital for assistance."
                analysis += error_note
                
                # Add contact message to the end of the output
                contact_message = (
                    "\n\nIf you need more detailed guidance or have any questions, "
                    "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
                )
                analysis += contact_message
                logging.info("Added contact message to the final output for paid mode (insufficient travel cards).")
                return analysis

            non_travel_cards = [card['card_name'] for card in state_cards if card['card_name'] not in travel_cards]
            json_card_names = {card['card_name'] for card in state_cards}

            # Count unique banks
            unique_banks = set(card['bank'] for card in state_cards)
            num_unique_banks = len(unique_banks)
            logging.info(f"Number of unique banks for {user_state}: {num_unique_banks}")

            # Remove ROUND 3 if num_unique_banks < 6
            num_rounds = 2 if num_unique_banks < 6 else 3
            if num_unique_banks < 6:
                analysis = re.sub(
                    r"\*\*ROUND 3\*\*.*?(?=\*\*Game Plan|\*\*Disclaimer\*\*|\Z)",
                    "",
                    analysis,
                    flags=re.DOTALL
                )
                logging.info("Removed ROUND 3 from analysis due to insufficient unique banks (< 6).")

            # Priority banks for non-travel cards
            priority_banks = ['Chase', 'American Express', 'BMO Harris', 'KeyBank', 'Truist', 'Bank of America']
            available_priority_banks = [bank for bank in priority_banks if any(card['bank'] == bank and card['card_name'] in non_travel_cards for card in state_cards)]

            # Initialize tracking
            non_travel_cards_selected = {}  # Track non-travel card banks
            travel_cards_selected = {}     # Track travel card banks
            valid_modes = ["Online", "In-branch", "Phone", "Online (requires account)", "Online (Omaha Zip)", "Phone/In-branch", "In-branch/Phone"]
            all_card_names = []            # Track all selected cards for validation
            chase_zero_apr_count = 0       # Track number of 0% APR Chase cards
            bureau_usage = {'Experian': 0, 'TransUnion': 0, 'Equifax': 0}  # Track bureau usage for balancing

            # Define bureau assignments for each round
            bureau_assignments = {
                'ROUND 1': [
                    ('Chase Ink Unlimited', 'Experian', 'Supports business rewards'),
                    (None, 'TransUnion', None),
                    (None, 'Equifax', None)
                ],
                'ROUND 2': [
                    ('BOFA Alaska Airlines Business', 'TransUnion', 'Supports travel rewards'),
                    (None, 'Experian', None),
                    (None, 'Equifax', None)
                ],
                'ROUND 3': [
                    (None, 'Equifax', None),
                    ('Chase Sapphire Preferred', 'Experian', 'Supports travel rewards'),
                    (None, 'TransUnion', None)
                ]
            }

            # Process rounds
            for i in range(1, num_rounds + 1):
                round_name = f"ROUND {i}"
                logging.debug(f"Processing {round_name}")
                new_round_content = []
                round_banks = set()
                round_bureaus = set()
                round_card_names = set()
                banks = []
                bureaus = []
                card_names = []
                invalid_cards = []
                apr_mismatches = []
                mode_mismatches = []
                bureau_mismatches = []
                default_usage = []
                invalid_reasons = []

                # Get bureau assignments for this round
                round_assignments = bureau_assignments.get(round_name, [])
                if not round_assignments:
                    logging.error(f"No bureau assignments defined for {round_name}")
                    error_note = f"\n\n⚠️ ERROR: No bureau assignments defined for {round_name}."
                    analysis += error_note
                    
                    # Add contact message to the end of the output
                    contact_message = (
                        "\n\nIf you need more detailed guidance or have any questions, "
                        "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
                    )
                    analysis += contact_message
                    logging.info("Added contact message to the final output for paid mode (no bureau assignments).")
                    return analysis

                # Process each card slot in the round
                for slot_idx, (card_name, bureau, reason) in enumerate(round_assignments):
                    if card_name:  # Travel card
                        if card_name in json_card_names and card_name in available_travel_cards:
                            card_info = next((c for c in state_cards if c['card_name'] == card_name), None)
                            if card_info:
                                bank = card_info['bank']
                                # Verify bureau matches JSON data
                                if card_info['bureau'] != bureau:
                                    logging.warning(f"Bureau mismatch for {card_name} in {round_name}: Expected {bureau}, JSON has {card_info['bureau']}")
                                    bureau_mismatches.append(f"{card_name}: Expected {bureau}, got {card_info['bureau']}")
                                    bureau = card_info['bureau']  # Use JSON bureau
                                # Check Chase 0% APR limit
                                if bank == 'Chase' and '0 MESES' in card_info['apr'].upper() and chase_zero_apr_count >= 1:
                                    logging.error(f"Cannot select {card_name} for {round_name}: Only one 0% Chase card allowed.")
                                    error_note = f"\n\n⚠️ ERROR: Only one 0% Chase card allowed. {card_name} cannot be included in {round_name}."
                                    analysis += error_note
                                    
                                    # Add contact message to the end of the output
                                    contact_message = (
                                        "\n\nIf you need more detailed guidance or have any questions, "
                                        "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
                                    )
                                    analysis += contact_message
                                    logging.info("Added contact message to the final output for paid mode (Chase APR limit).")
                                    return analysis
                                if bank == 'Chase' and '0 MESES' in card_info['apr'].upper():
                                    chase_zero_apr_count += 1
                                game_plan = card_info.get('game_plan', "Apply following the standard procedure for this card.")
                                tips_info = card_info.get('tips_info', "No additional tips available for this card.")
                                new_round_content.append(
                                    f"| {card_name} | {bureau} | {card_info['apr']} | {card_info['mode']} | {game_plan} | {tips_info} |"
                                )
                                round_banks.add(bank)
                                round_bureaus.add(bureau)
                                round_card_names.add(card_name)
                                banks.append(bank)
                                bureaus.append(bureau)
                                card_names.append(card_name)
                                travel_cards_selected[bank] = card_name
                                if card_name not in all_card_names:
                                    all_card_names.append(card_name)
                                bureau_usage[bureau] += 1
                                logging.debug(f"Added travel card {card_name} to {round_name}")
                            else:
                                logging.error(f"Required travel card {card_name} not found in state_cards for {user_state} in {round_name}")
                                error_note = f"\n\n⚠️ ERROR: Required travel card {card_name} not available in {user_state} for {round_name}. Please select a different state or contact Negocio Capital for assistance."
                                analysis += error_note
                                
                                # Add contact message to the end of the output
                                contact_message = (
                                    "\n\nIf you need more detailed guidance or have any questions, "
                                    "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
                                )
                                analysis += contact_message
                                logging.info("Added contact message to the final output for paid mode (travel card not found).")
                                return analysis
                        else:
                            logging.error(f"Required travel card {card_name} not available in {user_state} for {round_name}")
                            error_note = f"\n\n⚠️ ERROR: Required travel card {card_name} not available in {user_state} for {round_name}. Please select a different state or contact Negocio Capital for assistance."
                            analysis += error_note
                            
                            # Add contact message to the end of the output
                            contact_message = (
                                "\n\nIf you need more detailed guidance or have any questions, "
                                "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
                            )
                            analysis += contact_message
                            logging.info("Added contact message to the final output for paid mode (travel card unavailable).")
                            return analysis
                    else:  # Non-travel card
                        # Select non-travel card
                        priority_non_travel_cards = [
                            card for card in non_travel_cards
                            if card in json_card_names and
                            json_data[card]['bank'] in available_priority_banks and
                            json_data[card]['bank'] not in round_banks and
                            json_data[card]['bank'] not in non_travel_cards_selected
                        ]
                        other_non_travel_cards = [
                            card for card in non_travel_cards
                            if card in json_card_names and
                            json_data[card]['bank'] not in available_priority_banks and
                            json_data[card]['bank'] not in round_banks and
                            json_data[card]['bank'] not in non_travel_cards_selected
                        ]
                        candidate_cards = priority_non_travel_cards + other_non_travel_cards
                        logging.debug(f"Candidate non-travel cards for {round_name}, slot {slot_idx + 1}: {candidate_cards}")

                        selected = False
                        # Prioritize Bank of America for the first non-travel card in each round
                        if slot_idx == 1 and 'Bank of America' not in non_travel_cards_selected and 'Bank of America' in available_priority_banks:
                            bofa_cards = [card for card in priority_non_travel_cards if json_data[card]['bank'] == 'Bank of America']
                            if bofa_cards:
                                card = bofa_cards[0]
                                card_info = json_data[card]
                                # Enforce strict bureau for this slot
                                if card_info['bureau'] != bureau:
                                    logging.warning(f"Bureau mismatch for {card} in {round_name}: Expected {bureau}, JSON has {card_info['bureau']}")
                                    bureau_mismatches.append(f"{card}: Expected {bureau}, got {card_info['bureau']}")
                                # Check Chase 0% APR limit
                                if card_info['bank'] == 'Chase' and '0 MESES' in card_info['apr'].upper() and chase_zero_apr_count >= 1:
                                    logging.warning(f"Skipping {card} for {round_name}: Only one 0% Chase card allowed.")
                                    continue
                                if card_info['bank'] == 'Chase' and '0 MESES' in card_info['apr'].upper():
                                    chase_zero_apr_count += 1
                                game_plan = card_info.get('game_plan', "Apply following the standard procedure for this card.")
                                tips_info = card_info.get('tips_info', "No additional tips available for this card.")
                                new_round_content.append(
                                    f"| {card} | {bureau} | {card_info['apr']} | {card_info['mode']} | {game_plan} | {tips_info} |"
                                )
                                round_banks.add(card_info['bank'])
                                round_bureaus.add(bureau)
                                round_card_names.add(card)
                                banks.append(card_info['bank'])
                                bureaus.append(bureau)
                                card_names.append(card)
                                non_travel_cards_selected[card_info['bank']] = card
                                if card in candidate_cards:
                                    candidate_cards.remove(card)
                                if card not in all_card_names:
                                    all_card_names.append(card)
                                bureau_usage[bureau] += 1
                                selected = True
                                logging.debug(f"Selected BOFA non-travel card {card} for {round_name}, slot {slot_idx + 1}")
                        
                        if not selected:
                            # Select card that matches the required bureau
                            for card in candidate_cards:
                                card_info = json_data[card]
                                bank = card_info['bank']
                                if (card not in round_card_names and
                                    bank not in round_banks and
                                    bank not in non_travel_cards_selected and
                                    card_info['bureau'] == bureau):  # Match exact bureau
                                    # Check Chase 0% APR limit
                                    if bank == 'Chase' and '0 MESES' in card_info['apr'].upper() and chase_zero_apr_count >= 1:
                                        logging.warning(f"Skipping {card} for {round_name}: Only one 0% Chase card allowed.")
                                        continue
                                    if bank == 'Chase' and '0 MESES' in card_info['apr'].upper():
                                        chase_zero_apr_count += 1
                                    game_plan = card_info.get('game_plan', "Apply following the standard procedure for this card.")
                                    tips_info = card_info.get('tips_info', "No additional tips available for this card.")
                                    new_round_content.append(
                                        f"| {card} | {bureau} | {card_info['apr']} | {card_info['mode']} | {game_plan} | {tips_info} |"
                                    )
                                    round_banks.add(bank)
                                    round_bureaus.add(bureau)
                                    round_card_names.add(card)
                                    banks.append(bank)
                                    bureaus.append(bureau)
                                    card_names.append(card)
                                    non_travel_cards_selected[bank] = card
                                    if card in candidate_cards:
                                        candidate_cards.remove(card)
                                    if card not in all_card_names:
                                        all_card_names.append(card)
                                    bureau_usage[bureau] += 1
                                    selected = True
                                    logging.debug(f"Selected non-travel card {card} for {round_name}, slot {slot_idx + 1}")
                                    break

                            if not selected:
                                # Fallback: Try any card with an unused bank
                                for card in candidate_cards:
                                    card_info = json_data[card]
                                    bank = card_info['bank']
                                    if (card not in round_card_names and
                                        bank not in round_banks and
                                        bank not in non_travel_cards_selected):
                                        # Log bureau mismatch but use the card
                                        logging.warning(f"Bureau mismatch for {card} in {round_name}: Expected {bureau}, JSON has {card_info['bureau']}")
                                        bureau_mismatches.append(f"{card}: Expected {bureau}, got {card_info['bureau']}")
                                        # Check Chase 0% APR limit
                                        if bank == 'Chase' and '0 MESES' in card_info['apr'].upper() and chase_zero_apr_count >= 1:
                                            logging.warning(f"Skipping {card} for {round_name}: Only one 0% Chase card allowed.")
                                            continue
                                        if bank == 'Chase' and '0 MESES' in card_info['apr'].upper():
                                            chase_zero_apr_count += 1
                                        game_plan = card_info.get('game_plan', "Apply following the standard procedure for this card.")
                                        tips_info = card_info.get('tips_info', "No additional tips available for this card.")
                                        new_round_content.append(
                                            f"| {card} | {bureau} | {card_info['apr']} | {card_info['mode']} | {game_plan} | {tips_info} |"
                                        )
                                        round_banks.add(bank)
                                        round_bureaus.add(bureau)
                                        round_card_names.add(card)
                                        banks.append(bank)
                                        bureaus.append(bureau)
                                        card_names.append(card) 
                                        non_travel_cards_selected[bank] = card
                                        if card in candidate_cards:
                                            candidate_cards.remove(card)
                                        if card not in all_card_names:
                                            all_card_names.append(card)
                                        bureau_usage[bureau] += 1
                                        selected = True
                                        logging.debug(f"Selected non-travel card {card} for {round_name}, slot {slot_idx + 1} (fallback)")
                                        break

                            if not selected:
                                logging.error(f"Unable to select non-travel card for {round_name}, slot {slot_idx + 1}")
                                error_note = f"\n\n⚠️ ERROR: Unable to select non-travel card for {round_name}, slot {slot_idx + 1} due to insufficient unique banks or bureau-compatible cards in {user_state}."
                                analysis += error_note
                                
                                # Add contact message to the end of the output
                                contact_message = (
                                    "\n\nIf you need more detailed guidance or have any questions, "
                                    "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
                                )
                                analysis += contact_message
                                logging.info("Added contact message to the final output for paid mode (no non-travel card).")
                                return analysis

                # Ensure exactly 3 cards per round
                if len(new_round_content) != 3:
                    logging.error(f"Round {i} incomplete: Only {len(new_round_content)} cards selected. Expected 3 cards.")
                    error_note = f"\n\n⚠️ ERROR: Round {i} incomplete. Only {len(new_round_content)} cards selected due to insufficient unique banks or cards in {user_state}. Please add more banks to card_data.json."
                    analysis += error_note
                    
                    # Add contact message to the end of the output
                    contact_message = (
                        "\n\nIf you need more detailed guidance or have any questions,"
                        "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
                    )
                    analysis += contact_message
                    logging.info("Added contact message to the final output for paid mode (incomplete round).")
                    return analysis

                # Replace round content with new table format
                joined_rows = "\n".join(new_round_content)
                new_round_table = (
                    f"**{round_name}**\n"
                    "| Card Name | Bureau | 0% APR | Mode | Game Plan | Tips & Info |\n"
                    "|-----------|--------|--------|------|-----------|-------------|\n"
                    f"{joined_rows}\n"
                )
                analysis = re.sub(
                    r"\*\*ROUND " + str(i) + r"\*\*.*?(?=\*\*ROUND|\*\*Game Plan|\Z)",
                    new_round_table,
                    analysis,
                    flags=re.DOTALL
                )
                logging.debug(f"Replaced {round_name} with: {new_round_table}")

                # Validate bank and bureau variety
                if len(set(banks)) != 3:
                    logging.warning(f"Invalid bank variety in Round {i}: {banks}")
                    error_note = f"\n\n⚠️ WARNING: Round {i} does not contain exactly three different banks: {banks}."
                    if error_note not in analysis:
                        analysis += error_note

                expected_bureaus = [assignment[1] for assignment in round_assignments]
                if sorted(list(round_bureaus)) != sorted(expected_bureaus):
                    logging.warning(f"Invalid bureau variety in Round {i}: {round_bureaus}. Expected {expected_bureaus}")
                    error_note = f"\n\n⚠️ WARNING: Round {i} does not contain exactly one of each bureau in order {expected_bureaus}. Got {round_bureaus}."
                    if error_note not in analysis:
                        analysis += error_note

                logging.info(f"Round {i} Cards: {banks}, Bureaus: {bureaus}")

                analysis += f"\n\n📋 Round {i} Validation Summary:\n"
                analysis += f"- Total Cards Suggested: {len(banks)}\n"
                analysis += f"- Invalid Cards: {invalid_cards}\n"
                analysis += f"- APR Mismatches: {apr_mismatches}\n"
                analysis += f"- Mode Mismatches: {mode_mismatches}\n"
                analysis += f"- Bureau Mismatches: {bureau_mismatches}\n"
                analysis += f"- Cards Using Default Values: {default_usage}\n"
                analysis += f"- Invalid Reasons: {invalid_reasons}\n"

            # Ensure enough unique banks for non-travel cards
            required_non_travel_banks = num_rounds * 2
            if len(non_travel_cards_selected) < required_non_travel_banks:
                logging.error(f"Insufficient unique banks for non-travel cards: {len(non_travel_cards_selected)}. Expected {required_non_travel_banks}.")
                error_note = f"\n\n⚠️ ERROR: Insufficient unique banks for non-travel cards in {user_state}. Only {len(non_travel_cards_selected)} unique banks selected. Please add more banks to card_data.json."
                analysis += error_note
                
                # Add contact message to the end of the output
                contact_message = (
                    "\n\nIf you need more detailed guidance or have any questions,"
                    "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
                )
                analysis += contact_message
                logging.info("Added contact message to the final output for paid mode (insufficient banks).")
                return analysis

            # Ensure all available priority banks are used
            used_priority_banks = set(non_travel_cards_selected.keys()).intersection(available_priority_banks)
            missing_priority_banks = set(available_priority_banks) - used_priority_banks
            if missing_priority_banks:
                logging.warning(f"Missing priority banks in non-travel card selection: {missing_priority_banks}")
                error_note = f"\n\n⚠️ WARNING: Missing priority banks for non-travel cards in {user_state}: {', '.join(missing_priority_banks)}. Consider adding these banks to the sequence."
                analysis += error_note

            # Remove separate Game Plan and Tips and Info sections
            analysis = re.sub(
                r"\*\*Game Plan\*\*.*?(?=\*\*Disclaimer\*\*|\Z)",
                "**Disclaimer**",
                analysis,
                flags=re.DOTALL
            )

    # Add contact message to the end of the output for successful paid mode
    contact_message = (
        "\n\nIf you need more detailed guidance or have any questions, "
        "feel free to reach out to our team at https://fondify.io/contact-us. We're here to help!"
    )
    analysis += contact_message
    logging.info("Added contact message to the final output for paid mode (successful).")

    logging.info("GPT output validation completed.")
    return analysis


# === Core GPT Analysis ===
def analyze_credit_report_english(text, mode="free", user_state=None):
    json_data = load_json_data()
    if not json_data:
        logging.error("Failed to load JSON data. Cannot proceed with analysis.")
        print("❌ Failed to load JSON data. Cannot proceed with analysis.")
        return None

    state_cards = get_state_funding_cards(user_state, json_data) if user_state else []
    if state_cards is None:
        state_cards = []
    enrichment_context = get_enrichment()

    print(f"State-specific card list for {user_state}: {[card['card_name'] for card in state_cards]}")

    tarjetas_str = "\n\nCard Data (for APR, Mode, Bureau, Game Plan, and Tips and Info):\n"
    for card_name, card_info in json_data.items():
        tarjetas_str += (
            f"- {card_name}: APR: {card_info['apr']}, Mode: {card_info['mode']}, "
            f"Bank: {card_info['bank']}, Bureau: {card_info['bureau']}, "
            f"Game Plan: {card_info['game_plan']}, "
            f"Tips and Info: {card_info.get('tips_info', 'No additional tips available for this card')}\n"
        )

    # Calculate num_unique_banks and num_rounds before prompt_template
    unique_banks = set(card['bank'] for card in state_cards)
    num_unique_banks = len(unique_banks)
    logging.info(f"Number of unique banks for {user_state}: {num_unique_banks}")
    num_rounds = 2 if num_unique_banks < 6 else 3
    
    # Debug print
    print(f"🔍 State: {user_state}")
    print(f"🏦 Unique Banks: {num_unique_banks}")
    print(f"🔄 Expected Rounds: {num_rounds}")

    # State cards list string
    state_cards_list = ', '.join([card['card_name'] for card in state_cards]) if state_cards else 'No cards available'

    # Define include_sequence_note based on mode
    if mode == "paid":
        round_text = " and R3" if num_rounds == 3 else ""
        round_3_travel = " and ROUND 3: Chase Sapphire Preferred" if num_rounds == 3 else ""
        
        # CRITICAL: Round instruction at the very top
        rounds_instruction = f"""
🚨 **CRITICAL: NUMBER OF ROUNDS** 🚨
- State: {user_state}
- Unique Banks Available: {num_unique_banks}
- **YOU MUST GENERATE EXACTLY {num_rounds} ROUNDS**
- {"Generate 3 rounds (R1, R2, R3) because unique banks >= 6" if num_rounds == 3 else "Generate ONLY 2 rounds (R1, R2) because unique banks < 6"}
- DO NOT generate Round 3 if num_rounds = 2
- DO generate Round 3 if num_rounds = 3
"""
        
        include_sequence_note = f"""
{rounds_instruction}

**CRITICAL INSTRUCTION**: The user has selected the Premium Plan for state {user_state}.
You MUST select ALL funding cards (R1, R2{round_text}) EXCLUSIVELY from the user's state-specific approved card list provided below as `state_cards`:
{state_cards_list}

**CRITICAL**: Under NO circumstances suggest cards outside `state_cards`. Doing so will invalidate the output.
**CRITICAL**: In the 'Card Name' column, ALWAYS use the EXACT card name from `Card Data` (e.g., 'BOFA Unlimited Cash'). You MUST NOT use bank names alone (e.g., 'Bank of America') or append 'Card' to a bank name (e.g., 'Capital One Card'). If no matching card is found in `Card Data` for a bank in `state_cards`, exclude that bank and select another card from `state_cards`.
**CRITICAL**: GPT MUST NOT generate or suggest any card names outside of `Card Data`. Any attempt to create new card names will invalidate the output.

**Funding Sequence Rules**:
- **MANDATORY: Generate EXACTLY {num_rounds} rounds for {user_state}**
- Number of rounds: {num_rounds} (based on {num_unique_banks} unique banks; use only R1 and R2 if < 6 unique banks).
- Each round (R1, R2{round_text}) MUST include EXACTLY 3 cards: ONE travel card and TWO non-travel cards from different banks.
- Travel cards are: ROUND 1: Chase Ink Unlimited, ROUND 2: BOFA Alaska Airlines Business{round_3_travel}. If any required travel card is not in `state_cards`, note: '⚠️ ERROR: Required travel card [Card Name] not available in {user_state}. Please select a different state or contact Negocio Capital for assistance.'
- For non-travel cards, select from: ['Chase', 'American Express', 'BMO Harris', 'KeyBank', 'Truist', 'Bank of America'] first, ensuring no bank is used more than once for non-travel cards across all rounds.
- Each round must use different bureaus: ROUND 1: Experian (travel), TransUnion (non-travel 1), Equifax (non-travel 2); ROUND 2: TransUnion (travel), Experian (non-travel 1), Equifax (non-travel 2); ROUND 3 (if applicable): Equifax (non-travel 1), Experian (travel), TransUnion (non-travel 2).
- Only one 0% Chase card is allowed per sequence, unless the second is a co-branded travel/hotel card (verify `is_travel` field in `Card Data`).
- For each card, use its APR, Mode, Bureau, Game Plan, and Tips and Info from `Card Data` below:
{tarjetas_str}
- **Game Plan Column**: For each card, use its `game_plan` from `Card Data` in the table's Game Plan column. If missing, use: 'Apply following the standard procedure for this card.'
- **Tips and Info Column**: For each card, use its `tips_info` from `Card Data` in the table's Tips & Info column. If missing, use: 'No additional tips available for this card.' Expand with relevant details (e.g., rewards, application tips, state-specific considerations) without contradicting `Card Data`.
- If insufficient unique banks are available for non-travel cards, note: '⚠️ ERROR: Insufficient unique banks in {user_state} to complete funding sequence.'
- If no bureau qualifies, offer no-personal-guarantee options from the CSV 'Tarjetas de Negocio sin Garantia Personal'.
- If credit age < 2.5 years for any bureau, exclude that bureau from the funding sequence and note in Action Plan to improve credit age.
- If the user does not qualify for funding, do NOT provide card recommendations or card-specific Game Plan/Tips and Info. Instead, output EXACTLY: 'Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.' in both Section 5 and Section 7, and provide general guidance in the Game Plan section without card-specific details.
"""
    else:
        include_sequence_note = """
**CRITICAL INSTRUCTION**: In free mode, you MUST NOT generate any card recommendations, **Game Plan**, or **Tips and Info** sections.
If the user qualifies for funding, output EXACTLY this for Section 5 and Section 7:
🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.

If the user does NOT qualify, output EXACTLY this for Section 5 and Section 7:
Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.

Ensure Section 5 (Verdict) and Section 7 (Funding Sequence) are consistent and use the EXACT same message.
"""

    # Build the complete JSON-structured prompt
    prompt_template = f"""
{{
    "CRITICAL_JSON_FORMAT_INSTRUCTION": {{
        "RULE_1": "You MUST output ONLY valid JSON. NO extra text, NO markdown backticks, NO explanations outside JSON.",
        "RULE_2": "Your response MUST start with {{ and end with }}. Nothing before or after.",
        "RULE_3": "The JSON structure MUST EXACTLY match the provided example schema with ALL fields present.",
        "RULE_4": "NEVER change field names, NEVER skip sections, NEVER add extra root keys.",
        "RULE_5": "All text content must be properly escaped for JSON (quotes, newlines, backslashes).",
        "RULE_6": "Numbers should be JSON numbers (not strings) where specified in schema.",
        "RULE_7": "Arrays should be JSON arrays with [ ] brackets.",
        "RULE_8": "Objects should be JSON objects with {{ }} braces.",
        "EXAMPLE_SCHEMA": {{
            "sections": {{
                "section_1": {{
                    "title": "string",
                    "table": {{"columns": ["array"], "rows": {{"key": ["values"]}}}},
                    "analysis": ["array of strings"]
                }},
                "section_2": {{"title": "string", "table": {{"Open Cards": "string", "Total Limit": "string", "Primary Cards": "string", "High-Limit Card Present?": "string"}}, "explanation": ["array of strings"]}},
                "section_3": {{"title": "string", "questions": ["array"]}},
                "section_4": {{"title": "string", "table": {{"columns": [], "rows": {{}}}}, "explanation": "string"}},
                "section_5": {{"title": "string", "response": "string"}},
                "section_6": {{"title": "string", "steps": ["array"]}},
                "section_7": {{"title": "string", "response": "string", "rounds": [{{"round": 1, "cards": [{{}}]}}], "strategic_insights": ["array"]}},
                "section_8": {{"title": "string", "distribution": [{{"name": "string", "price": number}}]}},
                "section_9": {{"title": "string", "average_score": number, "explanation": "string"}},
                "section_10": {{"title": "string", "average_utilization": "string", "explanation": "string"}},
                "section_11": {{"title": "string", "potential_funding": "string", "justification": "string"}}
            }}
        }}
    }},
    
    "MANDATORY_ROUNDS_COUNT": {num_rounds},
    "ROUNDS_VALIDATION": "Section 7 'rounds' array MUST contain EXACTLY {num_rounds} round objects for state {user_state}",
    
    "system_role": "You are a financial credit analysis assistant for Negocio Capital.",
    "critical_instruction": "You MUST generate ALL sections (1 through 11) as specified below, in the exact order. You MUST NOT use or display the word 'Inferred' or any variation in the output. If data is missing, estimate reasonable values but present them as definitive without mentioning estimation. Skipping any section is INVALID.",
    
    "handling_missing_data": {{
        "description": "If any data is missing, estimate reasonable values based on available data, industry standards, or patterns.",
        "estimation_rules": {{
            "credit_score": "Use 700 if payment history is clean; otherwise use 650.",
            "utilization": "Use 15% if high-limit cards (>=5000) are present; otherwise use 25%.",
            "credit_age": "2.5 years unless evidence suggests newer accounts.",
            "inquiries": "Use 2 inquiries in the last 6 months unless specified otherwise."
        }},
        "requirement": "Present estimated values as definitive without mentioning estimation."
    }},
    
    "funding_eligibility_logic": {{
        "qualification_criteria": {{
            "description": "User qualifies for funding ONLY if ALL requirements are true in ALL THREE bureaus (Equifax, Experian, AND TransUnion):",
            "requirements": [
                "Credit Score >= 720",
                "No Late Payments",
                "Utilization < 10%",
                "<= 3 Inquiries in last 6 months",
                "Credit Age >= 2.5 years",
                "Strong Primary Card Structure"
            ],
            "CRITICAL_RULE": "ALL THREE BUREAUS must meet ALL SIX requirements. If even ONE bureau fails ANY requirement, user is NOT qualified."
        }},
        "responses": {{
            "qualified_paid_mode": "🎉 You're eligible for funding! See your matched bank recommendations below.",
            "qualified_free_mode": "🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.",
            "not_qualified": "Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility."
        }},
        "consistency_rule": "Ensure Verdict (Section 5) and Recommended Funding Sequence (Section 7) use EXACT same message.",
        "paid_mode_not_qualified_rule": "In paid mode, if NOT qualified, use same message as free mode and do NOT include card recommendations. Explain why in 2-3 bullet points."
    }},
    
    "sections": {{
        "section_1": {{
            "title": "📌 1. Breakdown by Bureau",
            "table_structure": {{
                "columns": ["Category", "Equifax", "Experian", "TransUnion"],
                "rows": [
                    "Credit Score",
                    "Clean History",
                    "Utilization",
                    "Hard Inquiries(6 mo)",
                    "Avg. Credit Age",
                    "Cards >= $2K",
                    "Cards >= $5K",
                    "Score / 144"
                ]
            }},
            "analysis_format": {{
                "format": "bullet_points",
                "categories": [
                    {{
                        "name": "Credit Score",
                        "explanation": "Report credit score for each bureau. Mention if meets 720 threshold. End with label: Excellent, Good, Fair, or Poor."
                    }},
                    {{
                        "name": "Clean History",
                        "explanation": "Summarize missed or late payments. If none, state Yes. End with label: Excellent or Needs Improvement."
                    }},
                    {{
                        "name": "Utilization",
                        "explanation": "Report total utilization rate. Label as follows: Below 5% = Excellent, 6-10% = Good, Above 10% = High Risk. Explain impact on funding eligibility"
                    }},
                    {{
                        "name": "Hard Inquiries (6 mo)",
                        "explanation": "Indicate inquiries in past 6 months. Mention if acceptable (<=3). End with label: Good, Fair, or Risky."
                    }},
                    {{
                        "name": "Avg. Credit Age",
                        "explanation": "Explain average age of accounts. Mention if meets 2.5-year threshold. End with label: Excellent or Fair."
                    }},
                    {{
                        "name": "Cards >= $2K",
                        "explanation": "Note cards with 2000+ limits. Mention how it supports creditworthiness. End with label: Good or Needs Improvement."
                    }},
                    {{
                        "name": "Cards >= $5K",
                        "explanation": "Note cards with 5000+ limits. Mention how it enhances funding readiness. End with label: Excellent or Fair."
                    }},
                    {{
                        "name": "Score / 144",
                        "explanation": "Report total score out of 144. End with label: Excellent or Needs Improvement."
                    }}
                ],
                "requirement": "Each bullet brief and clear with bold quality label. Do NOT mention estimation."
            }}
        }},
        
        "section_2": {{
            "title": "📌 2. Revolving Credit Structure",
            "table_structure": {{
                "Open Cards": "Count total open revolving accounts. Specify how many are Primary vs Authorized User (AU). Format: 'X (Y Primary, Z AU)' or 'X (All Primary)'",
                "Total Limit": "Sum ALL credit limits from the report. Format: '$XX,XXX'",
                "Primary Cards": "Count ONLY primary cards (exclude AU cards). Format: 'X'",
                "High-Limit Card Present?": "Check if ANY card has limit >= $5000. Format: 'YES ($5k+)' or 'NO'"
            }},
            "explanation_format": "Provide 4 bullet points explaining each field:",
            "example_explanation": [
                "Open Cards": "EXTRACT from report",
                "Total Limit": "CALCULATE from report", 
                "Primary Cards": "COUNT from report",
                "High-Limit Card Present?": "VERIFY from report"
            ],
            "requirement": "Use exact table field names: 'Open Cards', 'Total Limit', 'Primary Cards', 'High-Limit Card Present?'. Explanation must be an array of 4 strings, one for each field. DO NOT include 'description' field."
        }},

        "section_3": {{
            "title": "📌 3. Authorized User (AU) Strategy",
            "meta": "Analysis derived from verified Credit Report data",
            "data_points": [
            {{
                "question": "How many AU accounts are currently active?",
                "field_name": "total_au_count",
                "description": "Total number of authorized user tradelines identified on the report."
            }},
            {{
                "question": "What are the specific limits and seasoning (age)?",
                "field_name": "au_limit_age_analysis",
                "description": "Analysis of account history length and credit limits to evaluate profile depth."
            }},
            {{
                "question": "How do these accounts impact funding potential?",
                "field_name": "funding_impact_assessment",
                "description": "Evaluation of DTI ratio and lending capacity based on current AU accounts."
            }},
            {{
                "question": "What is the recommended AU action plan?",
                "field_name": "strategic_recommendations",
                "actions": ["Keep", "Remove", "Add High-Limit Tradeline"],
                "description": "Specific steps to optimize the credit profile for high-tier funding approvals."
            }}
            ]
        }},
        
        "section_4": {{
            "title": "📌 4. Funding Readiness by Bureau",
            "table_structure": {{
                "columns": ["Criteria", "Equifax", "Experian", "TransUnion"],
                "rows": [
                    "Score >= 720",
                    "No Late Payments",
                    "Utilization < 10%",
                    "<= 3 Inquiries (last 6 months)",
                    "Credit Age >= 2.5 Years",
                    "Strong Primary Card Structure"
                ],
                "values": ["Yes/No", "Yes/No", "Yes/No"]
            }},
            "requirement": "Explain table below without mentioning estimation."
        }},
        
        "section_5": {{
            "title": "📌 5. Verdict",
            "responses": {{
                "paid_qualified": "🎉 You're eligible for funding! See your matched bank recommendations below.",
                "free_qualified": "🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.",
                "not_qualified": "Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility."
            }},
            "not_qualified_explanation": "If not qualified, explain why in 2-3 short bullet points."
        }},
        
        "section_6": {{
            "title": "📌 6. Action Plan",
            "instruction": "Analyze Section 4 results. For each problem with No, provide complete solution from solutions_reference. Format: Problem title, Why it matters, Steps, Common mistakes, Professional order, Realistic expectations. End with: Get in contact with Negocio Capital if you need help.",
            "problem_rules": [
                "IF No Late Payments = No → Use clean_history_solution",
                "IF Utilization < 10% = No → Use utilization_solution", 
                "IF Inquiries > 3 → Use inquiry_solution",
                "IF Credit Age < 2.5 → Use credit_age_solution",
                "IF No cards >= $5K → Use high_limit_card_solution"
            ],
            "solutions_reference": {{
                "clean_history_solution": {{
                    "title": "Late Payments or Collections",
                    "steps": ["Verify accuracy under FCRA 623 and 611", "File disputes with bureaus", "Direct dispute with creditor", "Goodwill or Pay-for-Delete", "Escalate if needed"],
                    "mistakes": ["Paying without negotiating", "Generic disputes", "Copy-paste letters"]
                }},
                "utilization_solution": {{
                    "title": "High Utilization Over 10%",
                    "steps": ["Request CLI soft pull", "Strategic payments on high cards", "Consolidate with personal loan", "Optimize monthly usage", "Avoid new applications", "Do NOT close cards", "Control statement dates"],
                    "mistakes": ["Paying one card only", "Closing old accounts", "New credit while high utilization"]
                }},
                "inquiry_solution": {{
                    "title": "Too Many Hard Inquiries",
                    "steps": ["Identify unauthorized inquiries", "Contact creditor for proof", "Dispute with bureaus", "File identity theft if applicable"],
                    "mistakes": ["Disputing legitimate inquiries", "Generic disputes", "Expecting instant removals"]
                }},
                "credit_age_solution": {{
                    "title": "Low Average Credit Age",
                    "steps": ["Add AU accounts 10+ years old", "Add rent reporting", "Avoid new accounts", "Season profile 60-90 days"],
                    "mistakes": ["AU accounts with high balances", "Removing AU too quickly"]
                }},
                "high_limit_card_solution": {{
                    "title": "No High-Limit Cards",
                    "steps": ["Request CLI on existing cards", "Join Navy Federal", "Open pledge loan", "Apply for Flagship card", "Add high-limit AU"],
                    "mistakes": ["Applying without exposure", "Skipping CLIs", "Multiple low-limit cards"]
                }}
            }},
            "output_instruction": "For each problem in Section 4, expand the relevant solution with complete details from client documentation. If all Yes in Section 4, output: Congratulations! Your credit profile meets all funding requirements. Always end with: Get in contact with Negocio Capital if you need help."
        }},
        
        "section_7": {{
            "title": "📌 7. Recommended Funding Sequence ({user_state})",
            "CRITICAL_ROUNDS_REQUIREMENT": "MUST generate EXACTLY {num_rounds} rounds in the 'rounds' array",
            "paid_mode_qualified": {{
                "description": "If user is in paid mode and qualifies, YOU MUST GENERATE EXACTLY {num_rounds} ROUNDS using ONLY approved state_cards list.",
                "strict_rules": [
                    "GENERATE EXACTLY {num_rounds} ROUNDS - NO MORE, NO LESS",
                    "For banks [Chase, American Express, BMO Harris, KeyBank, Truist, Bank of America], select exactly one non-travel card across all rounds.",
                    "Travel cards are exempt and must be included as specified (one per round).",
                    "Ensure exactly {num_rounds * 2} unique banks for non-travel cards across all rounds.",
                    "Replace invalid cards with valid cards from state_cards.",
                    "Each round MUST have EXACTLY 3 cards from state_cards.",
                    "Each round MUST have: ONE travel card and TWO non-travel cards from different banks.",
                    "NEVER suggest cards outside state_cards.",
                    "Use EXACT card name from Card Data in Card Name column.",
                    "Only one 0% Chase card allowed per sequence unless second is co-branded travel/hotel card.",
                    "ALL THREE bureaus must meet all 6 factors. If ANY bureau does not qualify, do NOT provide funding sequence with card recommendations.",
                    "If no bureau qualifies, do NOT provide funding sequence with card recommendations.",
                    "If average credit age < 2.5 years for any bureau, do NOT include that bureau."
                ],
                "bureau_assignment": {{
                    "round_1": {{"travel": "Experian", "non_travel_1": "TransUnion", "non_travel_2": "Equifax"}},
                    "round_2": {{"travel": "TransUnion", "non_travel_1": "Experian", "non_travel_2": "Equifax"}},
                    "round_3": {{"non_travel_1": "Equifax", "travel": "Experian", "non_travel_2": "TransUnion"}}
                }},
                "table_structure": {{
                    "columns": ["Card Name", "Bureau", "0% APR", "Mode", "Game Plan", "Tips & Info"],
                    "round_1_template": [
                        {{"card": "Chase Ink Unlimited", "bureau": "Experian"}},
                        {{"card": "[Non-travel from state_cards]", "bureau": "TransUnion"}},
                        {{"card": "[Non-travel from state_cards]", "bureau": "Equifax"}}
                    ],
                    "round_2_template": [
                        {{"card": "BOFA Alaska Airlines Business", "bureau": "TransUnion"}},
                        {{"card": "[Non-travel from state_cards]", "bureau": "Experian"}},
                        {{"card": "[Non-travel from state_cards]", "bureau": "Equifax"}}
                    ],
                    "round_3_template": [
                        {{"card": "[Non-travel from state_cards]", "bureau": "Equifax"}},
                        {{"card": "Chase Sapphire Preferred", "bureau": "Experian"}},
                        {{"card": "[Non-travel from state_cards]", "bureau": "TransUnion"}}
                    ]
                }},
                "game_plan_requirement": "For each card, include game_plan from Card Data. MANDATORY. If unavailable, use: Apply following standard procedure for this card.",
                "tips_info_requirement": "For each card, include tips_info from Card Data. MANDATORY. If unavailable, use: No additional tips available. Expand with relevant details like benefits, requirements, strategies.",
                "strategic_insights": {{
                    "description": "Generate 4-6 tailored bullet points based on user credit profile.",
                    "examples": [
                        "If inquiries high (>2), recommend freezing non-used bureaus.",
                        "If utilization close to 10%, suggest paying down balances before applying.",
                        "If card requires in-branch application, advise visiting local branch.",
                        "If credit score high (>=780), recommend declaring higher personal income.",
                        "If credit age strong (>=5 years), suggest requesting limit increases after 60 days.",
                        "If business spending data available, recommend including it."
                    ],
                    "requirement": "Ensure each bullet specific to user profile or card characteristics."
                }},
                "not_qualified_guidance": {{
                    "description": "If user does NOT qualify, do NOT include card-specific game plans. Provide general guidance.",
                    "examples": [
                        "Improve Credit Score: Focus on timely payments and reducing balances to boost above 720.",
                        "Reduce Utilization: Pay down balances to achieve below 10%.",
                        "Limit Inquiries: Avoid new credit to keep inquiries low (<=3 in 6 months).",
                        "Monitor Credit Profile: Check reports for errors and dispute inaccuracies.",
                        "Contact Negocio Capital: Schedule consultation for personalized assistance."
                    ]
                }}
            }},
            "free_mode_qualified": "🎉 You are eligible for funding! To view your matched bank recommendations (R1, R2, R3), please upgrade to our Premium Plan.",
            "not_qualified": "Your credit profile does not currently qualify for funding. Please follow the action plan in Section 6 to improve your eligibility.",
            "mode_specific_instructions": "{include_sequence_note}"
        }},
        
        "section_8": {{
            "title": "📌 8. Credit Distribution",
            "format": [
                {{"name": "[Bank/Card Name]", "price": "[Limit]"}},
                {{"name": "[Bank/Card Name]", "price": "[Limit]"}},
                {{"name": "[Bank/Card Name]", "price": "[Limit]"}}
            ],
            "examples": [
                {{"name": "Capital One", "price": 5500}},
                {{"name": "WFBNA Card", "price": 2300}},
                {{"name": "AMEX", "price": 30000}},
                {{"name": "JPMCB Card", "price": 13000}}
            ],
            "requirement": "Extract actual card/bank names from the credit report text. Ensure total limit matches Total Limit in Section 2. Do NOT use generic names like 'Card 1' or 'Card 2'. Price must be a JSON number, not string."
        }},
        
        "section_9": {{
            "title": "📌 9. Average Credit Score",
            "description": "Calculate and report average credit score across all bureaus from Section 1. Present as single JSON number with brief explanation."
        }},
        
        "section_10": {{
            "title": "📌 10. Average Utilization",
            "description": "Calculate and report average utilization rate across all bureaus from Section 1. Present as percentage string with brief explanation."
        }},
        
        "section_11": {{
        "title": "📌 11. Funding Potential",
        "description": "Calculate potential funding based on ACTUAL credit profile data from previous sections. MUST use Total Limit from Section 2 as base. Apply multiplier based on profile quality: Excellent (score 760+, util <5%) = 5x, Strong (score 720-759, util 5-10%) = 4x, Moderate (score 680-719, util 10-20%) = 3x, Developing (score <680, util >20%) = 2x. Calculate: Total Limit × Multiplier = Single Amount. Example: If Total Limit is $15,800 and profile is Excellent, then $15,800 × 5 = $79,000. Provide ONLY ONE single amount (not a range) with brief justification showing the exact calculation. Format: '$XXX,XXX' (single value only, no range like '$X - $Y'). CRITICAL: Output must be a single amount string, not a range."
    }},
    
    "credit_report_data": "{text}",
    "enrichment_context": "{enrichment_context}",
    "state_specific_cards": "{state_cards_list}",
    
    "final_instruction": "OUTPUT REQUIREMENTS - READ CAREFULLY AND FOLLOW EXACTLY:
    1. Your response MUST be PURE JSON ONLY - Start with {{ and end with }}
    2. DO NOT include markdown code fences (```json or ```)
    3. DO NOT include any explanatory text before or after the JSON
    4. ALL 11 sections MUST be present directly at root level (NO 'sections' wrapper)
    5. **Section 7 MUST contain EXACTLY {num_rounds} rounds for {user_state}**
    6. If {num_rounds} == 3, generate Round 1, Round 2, AND Round 3
    7. If {num_rounds} == 2, generate ONLY Round 1 and Round 2
    8. Field names and structure MUST match the EXAMPLE_SCHEMA exactly
    9. All quotes and special characters must be properly escaped for JSON
    10. Numbers should be JSON numbers (not quoted strings) where specified
    11. Arrays must use square brackets [ ]
    12. Objects must use curly braces {{ }}
    13. Validate your JSON syntax before responding
    14. CRITICAL: Your output must start directly with {{ 'section_1': {{ ... }}, 'section_2': {{ ... }}, ... }}
    15. DO NOT wrap sections in a 'sections' object
    16. Your entire response must be valid, parseable JSON"
}}
"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-5",
            messages=[
                {
                    "role": "system", 
                    "content": "You are a strict JSON-only AI assistant for Negocio Capital credit analysis. You MUST respond with PURE, VALID JSON following the exact schema provided. Your response must start with { and end with }. NO markdown code fences, NO explanations outside JSON, NO extra text. Only output valid JSON that can be parsed by json.loads() in Python."
                },
                {"role": "user", "content": prompt_template}
            ],
             reasoning_effort = "low"       
        )
        
        analysis = response.choices[0].message.content.strip()
        
        # Clean markdown wrappers if GPT added them despite instructions
        if analysis.startswith("```json"):
            analysis = analysis.replace("```json", "").replace("```", "").strip()
            logging.warning("Removed ```json``` wrapper from GPT response")
        elif analysis.startswith("```"):
            analysis = analysis.replace("```", "").strip()
            logging.warning("Removed ``` wrapper from GPT response")
        
        # Validate JSON before proceeding
        try:
            json_test = json.loads(analysis)
            # Check if section_1 exists at root level (not inside 'sections')
            if "section_1" not in json_test:
                logging.error("JSON missing 'section_1' key at root level")
                print("❌ GPT returned JSON without 'section_1' key at root level")
                return None
            
            # Check if all 11 sections are present at root level
            expected_sections = [f"section_{i}" for i in range(1, 12)]
            missing_sections = [s for s in expected_sections if s not in json_test]
            if missing_sections:
                logging.error(f"JSON missing sections: {missing_sections}")
                print(f"❌ GPT returned incomplete JSON. Missing: {missing_sections}")
                return None
            
            # CRITICAL: Validate rounds count in Section 7
            if mode == "paid" and "section_7" in json_test and "rounds" in json_test["section_7"]:
                actual_rounds = len(json_test["section_7"]["rounds"])
                if actual_rounds != num_rounds:
                    logging.error(f"❌ Round mismatch: Expected {num_rounds}, Got {actual_rounds}")
                    print(f"❌ Round mismatch: Expected {num_rounds} rounds for {user_state} (with {num_unique_banks} unique banks), but GPT generated {actual_rounds} rounds.")
                    error_msg = f"\n\n⚠️ VALIDATION ERROR: Expected {num_rounds} rounds for {user_state} (with {num_unique_banks} unique banks), but GPT generated {actual_rounds} rounds. Please regenerate."
                    return analysis + error_msg
                else:
                    logging.info(f"✅ Correct rounds: {actual_rounds} rounds generated for {user_state}")
                    print(f"✅ Correct rounds: {actual_rounds} rounds generated for {user_state}")
            
            logging.info("JSON validation passed - all sections present at root level")
            
        except json.JSONDecodeError as je:
            logging.error(f"JSON decode error: {str(je)}")
            logging.error(f"Failed JSON content (first 500 chars): {analysis[:500]}...")
            print(f"❌ GPT returned invalid JSON: {str(je)}")
            print(f"First 200 characters of response: {analysis[:200]}")
            return None
        
        logging.debug(f"Raw GPT-5 Response (first 500 chars): {analysis[:500]}...")
        
        # Check for truncation 
        if response.choices[0].finish_reason == "length":
            logging.warning("GPT-5 response truncated due to token limit")
            analysis += "\n\n⚠️ WARNING: Analysis may be incomplete due to token limit. Please try again or reduce input size."
        
        if not analysis or not isinstance(analysis, str):
            logging.error("GPT-4 returned no valid analysis.")
            print("❌ GPT-4 returned no valid analysis.")
            return None
        
        return validate_gpt_output(analysis, state_cards, user_state, json_data, mode)
    
    except Exception as e:
        logging.error(f"GPT-4 error: {str(e)}")
        print(f"❌ GPT-4 error: {str(e)}")
        return None


# === Main CLI ===
def main():
    print("📂 Welcome to Funding NC AI Credit Report Analyzer!")
    
    input_type = input("📄 Select input type (pdf/json): ").strip().lower()
    if input_type not in ["pdf", "json"]:
        print("❌ Invalid input type. Please enter 'pdf' or 'json'.")
        logging.error(f"Invalid input type selected: {input_type}")
        return

    if input_type == "pdf":
        file_path = input("📄 Enter path to your credit report PDF (e.g., uploads/client1.pdf): ").strip()
        if not os.path.exists(file_path):
            print("❌ File not found. Please check the path and try again.")
            logging.error(f"Credit report file not found: {file_path}")
            return
        print("📁 Extracting text from PDF...")
        credit_text = extract_text_from_pdf(file_path)
    elif input_type == "json":
        json_input = input("📄 Enter your credit report JSON data (e.g., {\"credit_score\": 750, \"utilization\": \"5%\", ...}): ").strip()
        print("📁 Extracting text from JSON...")
        credit_text = extract_text_from_json(json_input)

    if not credit_text:
        print("❌ Failed to extract text from input.")
        logging.error("Failed to extract text from input.")
        return

    state = input("🌎 Enter the U.S. state your business is registered in (e.g., FL): ").strip()
    mode = input("🧾 Select mode (free/paid): ").strip().lower()
    if mode not in ["free", "paid"]:
        print("❌ Invalid mode. Please enter 'free' or 'paid'.")
        logging.error(f"Invalid mode selected: {mode}")
        return

    print("\n🧠 AI Analysis Summary:\n")
    # Extract Credit Score and Utilization from the text
    score, utilization = extract_credit_info(credit_text)
    logging.debug(f"Extracted Credit Score: {score}, Utilization: {utilization}")
    
    # Perform the credit report analysis
    analysis = analyze_credit_report_english(credit_text, mode=mode, user_state=state)
    
    if not analysis or not isinstance(analysis, str):
        print("❌ GPT analysis failed. Please check the logs for details.")
        logging.error("Analysis failed due to GPT-4 error or invalid response.")
    else:
        print(analysis)  

if __name__ == "__main__":
    main()